# -*- coding: utf-8 -*-
import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
import json
import streamlit as st
import pandas as pd
import os
from datetime import date, datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont, ImageOps
import requests
from io import BytesIO
# --- NUEVOS IMPORTS PARA GOOGLE SHEETS ---
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import numpy as np
import estructura as mod_estructura
import mod_reportes as mod_dashboard
import repvencimientos as mod_vencimientos
import repcumpleanos as mod_cumpleanos
import repvacaciones as mod_vacaciones
import reportegeneral as mod_reportegeneral
import gestor_evaluaciones as mod_gestor_evaluaciones
import mod_registro
import mod_nomina


st.set_page_config(page_title="Gestión Roosevelt", page_icon="🎓", layout="wide")

# ==========================================
# 1. CONFIGURACIÓN Y CONSTANTES
# ==========================================
DB = "DB_SISTEMA_GTH.xlsx"
F_N = "MG. ARTURO JAVIER GALINDO MARTINEZ"
F_C = "JEFE DE GESTIÓN DEL TALENTO HUMANO"

MOTIVOS_CESE = ["Término de contrato", "Renuncia", "Despido", "Mutuo acuerdo", "Fallecimiento", "Otros"]

COLUMNAS = {
    "PERSONAL": ["dni", "apellidos y nombres", "link"],
    "DATOS GENERALES": ["dni", "sede", "sexo", "apellidos y nombres", "dirección", "estado civil", "fecha de nacimiento", "edad"], 
    "DATOS FAMILIARES": ["parentesco", "apellidos y nombres", "dni", "fecha de nacimiento", "edad", "estudios", "telefono"],
    "EXP. LABORAL": ["dni", "tipo de experiencia", "lugar", "puesto", "fecha de inicio", "fecha de fin", "motivo de cese"],
    "FORM. ACADEMICA": ["dni", "tipo de estudio", "institución educativa", "mención (especialidad / carrera / etc)", "año", "estado", "horas académicas", "grado o título obtenido"],
    "INVESTIGACION": ["id", "dni", "tipo de registro", "enlace cti vitae", "codigo renacyt", "nivel renacyt", "titulo de publicacion", "base de datos", "nombre de revista", "cuartil", "año de publicacion", "doi o url", "nombre del proyecto", "entidad financiadora", "rol en el proyecto", "monto adjudicado", "estado del proyecto", "nombre del semillero", "resolucion", "rol en el semillero", "estado del semillero"],
    # NUEVAS COLUMNAS DE CONTRATOS APLICADAS:
    "CONTRATOS": ["dni", "cargo", "AREA", "f_inicio", "f_fin", "tipo de trabajador", "modalidad", "temporalidad", "tipo contrato", "estado", "LINK"],
    "VACACIONES": ["periodo", "fecha de inicio", "fecha de fin", "días generados", "dias gozados", "saldo", "link"],
    "OTROS BENEFICIOS": ["periodo", "tipo de beneficio", "link"],
    "MERITOS Y DEMERITOS": ["periodo", "merito o demerito", "motivo", "link"],
    "EVALUACION DEL DESEMPEÑO": ["periodo", "merito o demerito", "motivo", "link"],
    "LIQUIDACIONES": ["periodo", "firmo", "link"]
}

# ==========================================
# ---> NUEVA FUNCIÓN: CONVERTIR LINK DE DRIVE A IMAGEN DIRECTA <---
# ==========================================
def obtener_link_directo_drive(url):
    """Convierte un link de compartir de Google Drive en un link directo de imagen."""
    if not isinstance(url, str) or not url.strip():
        return None
    if "drive.google.com" in url and "/d/" in url:
        try:
            # Extrae el ID del archivo del link de Drive
            file_id = url.split("/d/")[1].split("/")[0]
            # Formato más estable para forzar la vista de la imagen
            return f"https://drive.google.com/uc?export=view&id={file_id}"
        except:
            return url
    return url
# ==========================================
# 2. FUNCIONES DE DATOS (VERSIÓN DEFINITIVA)
# ==========================================

SCOPE = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
SHEET_NAME = "DB_SISTEMA_GTH" 

def obtener_credenciales():
    if "google_json" in st.secrets:
        import json
        creds_dict = json.loads(st.secrets["google_json"])
        return ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, SCOPE)
    else:
        return ServiceAccountCredentials.from_json_keyfile_name("credenciales.json", SCOPE)

@st.cache_data(ttl=600)
def load_data():
    creds = obtener_credenciales()
    client = gspread.authorize(creds)
    spreadsheet = client.open(SHEET_NAME)
    worksheets = spreadsheet.worksheets()
    
    dfs = {}
    for worksheet in worksheets:
        try:
            data = worksheet.get_all_records()
            df = pd.DataFrame(data)
            if not df.empty:
                # Limpieza agresiva: quitamos espacios, tildes y guiones bajos
                df.columns = [str(c).strip().lower()
                              .replace('á', 'a').replace('é', 'e')
                              .replace('í', 'i').replace('ó', 'o')
                              .replace('ú', 'u').replace('_', ' ') 
                              for c in df.columns]
                
                # Arreglo especial para CONTRATOS (Evita el error f_inicio)
                if worksheet.title == "CONTRATOS":
                    # Mapeamos cualquier variante a 'f_inicio'
                    for col in df.columns:
                        if 'inicio' in col: df.rename(columns={col: 'f_inicio'}, inplace=True)
                        if 'termino' in col or 'fin' in col: df.rename(columns={col: 'f_fin'}, inplace=True)
                
                # Limpieza de DNI
                if "dni" in df.columns:
                    df["dni"] = df["dni"].astype(str).str.strip().str.replace(r'\.0$', '', regex=True).str.zfill(8)
                
            dfs[worksheet.title] = df
        except Exception as e:
            st.error(f"Error en {worksheet.title}: {e}")
    return dfs

def save_data(dfs):
    creds = obtener_credenciales()
    client = gspread.authorize(creds)
    sheet = client.open(SHEET_NAME)

    for h, df in dfs.items():
        worksheet = sheet.worksheet(h)
        df_s = df.copy()
        
        # --- EL ESCUDO ANTI-CLONACIÓN ---
        # Si el sistema detecta columnas repetidas (como "area" o "AREA"), las borra y deja solo una
        df_s = df_s.loc[:, ~df_s.columns.duplicated()]
        
        # --- BLINDAJE TOTAL CONTRA CELDAS VACÍAS (FIX PARA GOOGLE SHEETS) ---
        # 1. Rellenamos los nulos nativos de Pandas
        df_s = df_s.fillna("")
        
        # 2. Convertimos todo a texto puro
        df_s = df_s.astype(str)
        
        # 3. Limpiamos cualquier "fantasma" que Python haya dejado al convertir a texto
        fantasmas = ["nan", "NaN", "NaT", "nat", "None", "<NA>"]
        for fantasma in fantasmas:
            df_s = df_s.replace(fantasma, "")
            
        # Homologamos las columnas a mayúsculas
        df_s.columns = [str(c).upper() for c in df_s.columns]
        
        worksheet.clear()
        worksheet.update([df_s.columns.values.tolist()] + df_s.values.tolist())
    
    # Limpia la memoria automáticamente para que no tengas que darle F5 a cada rato
    st.cache_data.clear()

def get_consolidated_contracts(df_c):
    # Función inteligente para fusionar contratos consecutivos
    if df_c.empty: return df_c
    df_c = df_c.copy()
    df_c['f_inicio'] = pd.to_datetime(df_c['f_inicio'], errors='coerce')
    df_c['f_fin'] = pd.to_datetime(df_c['f_fin'], errors='coerce')
    df_c = df_c.sort_values('f_inicio').dropna(subset=['f_inicio'])
    
    merged = []
    for _, row in df_c.iterrows():
        if not merged:
            merged.append(row.to_dict())
        else:
            last = merged[-1]
            # Si la fecha de inicio del nuevo contrato es justo un día después del fin del anterior (o antes)
            if pd.notnull(last['f_fin']) and row['f_inicio'] <= last['f_fin'] + pd.Timedelta(days=1):
                # Ampliamos la fecha final
                last['f_fin'] = max(last['f_fin'], row['f_fin']) if pd.notnull(row['f_fin']) else row['f_fin']
                # Actualizamos al cargo más reciente
                last['cargo'] = row['cargo'] 
            else:
                merged.append(row.to_dict())
    return pd.DataFrame(merged)

def gen_word(nom, dni, df_c):
    doc = Document()
    section = doc.sections[0]
    section.page_height, section.page_width = Inches(11.69), Inches(8.27)
    section.top_margin, section.bottom_margin = Inches(1.6), Inches(1.2)

    if os.path.exists("header.png"):
        p_h = section.header.paragraphs[0]
        p_h.paragraph_format.left_indent = Inches(-1.0)
        p_h.add_run().add_picture("header.png", width=Inches(8.27))

    if os.path.exists("footer.png"):
        p_f = section.footer.paragraphs[0]
        p_f.paragraph_format.left_indent = Inches(-1.0)
        p_f.add_run().add_picture("footer.png", width=Inches(8.27))

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tit = p_tit.add_run("CERTIFICADO DE TRABAJO")
    r_tit.bold, r_tit.font.name, r_tit.font.size = True, 'Arial', Pt(18)

    doc.add_paragraph("\nLa oficina de Gestión de Talento Humano De La Universidad Privada De Huancayo “Franklin Roosevelt”, certifica que:").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p_inf = doc.add_paragraph()
    p_inf.add_run("El(la) TRABAJADOR(A) ").bold = False
    p_inf.add_run(nom).bold = True
    p_inf.add_run(f", identificado(a) con DNI N° {dni}, laboró bajo el siguiente detalle:")

    # Obtenemos los contratos fusionados automáticamente
    df_merged = get_consolidated_contracts(df_c)

    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    for i, h in enumerate(["CARGO", "FECHA INICIO", "FECHA FIN"]):
        celda = t.rows[0].cells[i]
        celda.text = h
        celda.paragraphs[0].runs[0].font.bold = True

    for _, fila in df_merged.iterrows():
        celdas = t.add_row().cells
        celdas[0].text = str(fila.get('cargo', ''))
        celdas[1].text = pd.to_datetime(fila['f_inicio']).strftime('%d/%m/%Y') if pd.notnull(fila['f_inicio']) else ""
        celdas[2].text = pd.to_datetime(fila['f_fin']).strftime('%d/%m/%Y') if pd.notnull(fila['f_fin']) else ""

    doc.add_paragraph("\nSe expide el presente a solicitud del interesado para los fines que considere convenientes.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph(f"\nHuancayo, {date.today().strftime('%d/%m/%Y')}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f.add_run("\n\n__________________________\n" + F_N + "\n" + F_C).bold = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
# ==============================================================================
# FUNCIÓN 2: GENERAR PAPELETA DE VACACIONES INDIVIDUAL (Word Duplicado A4)
# ==============================================================================
def gen_papeleta_vac(apellidos, nombres, dni_b, position, f_ingreso, period, start_d, end_d, days):
    template_path = "Template_Papeleta.docx"
    
    if not os.path.exists(template_path):
        st.error(f"⚠️ No se encontró la plantilla en: {template_path}. Por favor crea el archivo Word.")
        return None

    doc = Document(template_path)
    
    hoy = date.today()
    meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    txt_firma = f"Huancayo, {hoy.day} de {meses[hoy.month-1]} de {hoy.year}"

    fin_dt = pd.to_datetime(end_d, errors='coerce')
    if pd.notnull(fin_dt):
        retorno_dt = fin_dt + pd.Timedelta(days=1)
        if retorno_dt.weekday() == 6:  # Si cae Domingo (6), pasa a Lunes
            retorno_dt += pd.Timedelta(days=1)
        str_retorno = retorno_dt.strftime("%d/%m/%Y")
    else:
        str_retorno = ""

    replacements = {
        "{{APELLIDOS}}": str(apellidos).upper(),
        "{{NOMBRES}}": str(nombres).upper(),
        "{{DNI}}": str(dni_b),
        "{{CARGO}}": str(position).upper(),
        "{{F_INGRESO}}": f_ingreso.strftime("%d/%m/%Y") if isinstance(f_ingreso, (date, datetime)) else str(f_ingreso),
        "{{PERIODO}}": str(period),
        "{{F_INICIO}}": start_d.strftime("%d/%m/%Y") if isinstance(start_d, (date, datetime)) else str(start_d),
        "{{F_FIN}}": end_d.strftime("%d/%m/%Y") if isinstance(end_d, (date, datetime)) else str(end_d),
        "{{F_RETORNO}}": str_retorno,
        "{{DIAS}}": str(days),
        "{{FECHA_FIRMA}}": txt_firma
    }

    def replace_in_element(element, reps):
        for run in element.runs:
            for key, value in reps.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

    for paragraph in doc.paragraphs:
        replace_in_element(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_element(paragraph, replacements)

    docx_stream = BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    return docx_stream

# ==========================================
# 3. ESTILOS CSS
# ==========================================
st.markdown("""
<style>
    .stApp { background-color: #4a0000 !important; }
    [data-testid="stHeader"] { display: none !important; }
    
    .stApp p, .stMarkdown p { color: #FFFFFF; } 
    .stApp h1, .stApp h2, .stApp h3 { color: #FFD700 !important; }
    
    [data-testid="stSidebar"] { background-color: #4a0000 !important; }
    [data-testid="stSidebar"] h3 { color: #FFD700 !important; font-weight: bold !important; }
    [data-testid="stSidebar"] [data-testid="stImage"] { background-color: #FFF9C4 !important; border: 4px solid #FFD700 !important; border-radius: 15px !important; padding: 10px !important; }
    div[role="radiogroup"] label { background-color: transparent !important; }
    div[role="radiogroup"] label p { color: #FFFFFF !important; font-weight: bold !important; font-size: 16px !important; }
    
    /* ========================================= */
    /* BOTONES CON MEJOR CONTRASTE               */
    /* ========================================= */
    div.stButton > button, [data-testid="stFormSubmitButton"] > button { 
        background-color: #FFD700 !important; /* Amarillo Roosevelt */
        border: 2px solid #4a0000 !important; 
        border-radius: 10px !important; 
    }

    /* Forzamos el color Guinda en TODO el texto de CUALQUIER botón */
    div.stButton > button *, [data-testid="stFormSubmitButton"] > button *,
    div.stButton > button p, [data-testid="stFormSubmitButton"] > button p { 
        color: #4a0000 !important; 
        font-weight: bold !important; 
        font-size: 16px !important; 
    }

    div.stButton > button:hover, [data-testid="stFormSubmitButton"] > button:hover { 
        background-color: #ffffff !important; 
        border-color: #FFD700 !important; 
    }

   /* ========================================= */
   /* FONDOS Y CAJAS DE TEXTO                   */
   /* ========================================= */
    [data-testid="stExpander"] { 
        background-color: #FFF9C4 !important; 
        border: 2px solid #FFD700 !important; 
        border-radius: 10px !important; 
    }
    
    [data-testid="stExpander"] details { background-color: transparent !important; }
    [data-testid="stExpander"] summary { background-color: #FFD700 !important; padding: 10px !important; border-radius: 8px 8px 0 0 !important; }
    [data-testid="stExpander"] summary p { color: #4a0000 !important; font-weight: bold !important; font-size: 16px !important; }

    /* Damos fondo blanco y borde a las cajas donde se escribe para que resalten sobre el crema */
    [data-baseweb="input"], [data-baseweb="select"], [data-baseweb="textarea"] { 
        background-color: #FFFFFF !important; 
        border: 1px solid #4a0000 !important; 
        border-radius: 5px !important; 
    }
    
    /* El texto que tú escribes será negro */
    .stApp input, .stApp select, .stApp textarea, [data-baseweb="select"] span { 
        color: #000000 !important; 
        font-weight: bold !important; 
        -webkit-text-fill-color: #000000 !important;
    }

    /* Fix para los mensajes de advertencia (Ej: Activa la casilla) */
    [data-testid="stAlert"] { 
        background-color: #FFF9C4 !important; 
        border: 2px solid #FFD700 !important; 
        border-radius: 10px !important;
    }
    [data-testid="stAlert"] p, [data-testid="stAlert"] span, [data-testid="stAlert"] svg { 
        color: #4a0000 !important; 
        font-weight: bold !important; 
        font-size: 16px !important; 
    }

    /* ========================================= */
    /* TABLAS INTERACTIVAS                       */
    /* ========================================= */
    [data-testid="stDataEditor"], [data-testid="stTable"], .stTable { background-color: white !important; border-radius: 10px !important; overflow: hidden !important; }
    
    [data-testid="stDataEditor"] .react-grid-HeaderCell span { 
        color: #000000 !important; 
        font-weight: 900 !important; 
        font-size: 15px !important; 
        text-transform: uppercase !important; 
    }
    
    thead tr th { background-color: #FFF9C4 !important; color: #000000 !important; font-weight: bold !important; text-transform: uppercase !important; border: 1px solid #f0f0f0 !important; }
    
    /* ========================================= */
    /* SUBTÍTULOS (LABELS DE LOS FORMULARIOS)    */
    /* ========================================= */
    
    /* 1. Subtítulos generales (Login, Buscar, Alta Trabajador) en color DORADO */
    label p, label span, .stApp label p { 
        color: #FFD700 !important; 
        font-weight: bold !important; 
        font-size: 16px !important; 
    }
    
    /* 2. Subtítulos SOLO dentro de los recuadros desplegables (fondo crema) en color GUINDA */
    [data-testid="stExpander"] label p, [data-testid="stExpander"] label span { 
        color: #4a0000 !important; 
        font-weight: bold !important; 
    }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 4. LÓGICA DE DATOS Y SESIÓN
# ==========================================
if "rol" not in st.session_state: st.session_state.rol = None

if st.session_state.rol is None:
    st.markdown("<h3 style='text-align: center; color: #FFD700;'>¡Tu talento es importante! :)</h3>", unsafe_allow_html=True)

    col_logo1, col_logo2, col_logo3 = st.columns([1, 1.2, 1])
    with col_logo2:
        if os.path.exists("Logo_amarillo.png"): st.image("Logo_amarillo.png", use_container_width=False)

    col1, col2, col3 = st.columns([1, 1.5, 1])
    with col2:
        u = st.text_input("USUARIO").lower().strip()
        p = st.text_input("CONTRASEÑA", type="password")
        st.markdown('<p style="color:white; text-align:center; font-weight:bold; margin-top:15px;">Bienvenido (a) al sistema de gestión de datos de los colaboradores</p>', unsafe_allow_html=True)

        if st.button("INGRESAR"):
            if u == "admin": st.session_state.rol = "Admin"
            elif u == "supervisor" and p == "123": st.session_state.rol = "Supervisor"
            elif u == "lector" and p == "123": st.session_state.rol = "Lector"
            else: st.error("Credenciales incorrectas")

            if st.session_state.rol: st.rerun()

else:
    dfs = load_data()
    es_lector = st.session_state.rol == "Lector"

    with st.sidebar:
        st.markdown("<br>", unsafe_allow_html=True)
        col_logo_1, col_logo_2, col_logo_3 = st.columns([1, 2, 1]) 
        with col_logo_2:
            if os.path.exists("Logo_guindo.png"): st.image("Logo_guindo.png", use_container_width=False)
        st.markdown("<br>", unsafe_allow_html=True)

        # --- LÓGICA DE MENÚS INTELIGENTES ---
        if "menu_p" not in st.session_state: st.session_state.menu_p = "🔍 Consulta"
        if "menu_r" not in st.session_state: st.session_state.menu_r = None
        if "menu_activo" not in st.session_state: st.session_state.menu_activo = "🔍 Consulta"

        def click_menu_p():
            st.session_state.menu_activo = st.session_state.menu_p
            st.session_state.menu_r = None # Apaga los reportes

        def click_menu_r():
            if st.session_state.menu_r is not None:
                st.session_state.menu_activo = st.session_state.menu_r
                st.session_state.menu_p = None # Apaga el menú principal

        st.markdown("### 🛠️ MENÚ PRINCIPAL")
        st.radio("Menú Principal", ["🔍 Consulta", "➕ Registro", "📊 Nómina General", "🏢 Estructura", "📋 Evaluaciones", "📈 Dashboard Desempeño"], key="menu_p", on_change=click_menu_p, label_visibility="collapsed")
        
        st.markdown("<h3 style='color: #FFD700;'>📊 REPORTES</h3>", unsafe_allow_html=True)
        # Usamos index=None para que se pueda desmarcar sin textos extraños
        st.radio("Reportes", ["Reporte General", "Cumpleañeros", "Vacaciones", "Vencimientos"], key="menu_r", on_change=click_menu_r, index=None, label_visibility="collapsed")
        
        m = st.session_state.menu_activo

        st.markdown("---")
        if st.button("🚪 Cerrar Sesión", key="btn_logout"):
            st.session_state.rol = None
            st.rerun()

    # === SECCIÓN CONSULTA ===
    if m == "🔍 Consulta":
        st.markdown("<h2 style='color: #FFD700;'>Búsqueda de Colaborador</h2>", unsafe_allow_html=True)

        df_per_consulta = dfs["PERSONAL"].copy()
        
        df_per_consulta["dni_str"] = df_per_consulta.get("dni", pd.Series([""]*len(df_per_consulta))).astype(str).str.strip()
        apellidos_col = df_per_consulta.get("apellidos", pd.Series([""]*len(df_per_consulta))).fillna("").astype(str).str.strip()
        nombres_col = df_per_consulta.get("nombres", pd.Series([""]*len(df_per_consulta))).fillna("").astype(str).str.strip()
        
        df_per_consulta["nom_str"] = (apellidos_col + " " + nombres_col).str.strip()
        df_per_consulta["search_str"] = df_per_consulta["dni_str"] + " - " + df_per_consulta["nom_str"]
        
        opciones_buscador = [""] + [x for x in df_per_consulta["search_str"].tolist() if x != " - "]

        selected_search = st.selectbox("🔍 Escriba el DNI o Apellidos y Nombres:", opciones_buscador)

        if selected_search:
            dni_buscado = selected_search.split(" - ")[0].strip()
            
            fila_pers = df_per_consulta[df_per_consulta["dni_str"] == dni_buscado]
            if not fila_pers.empty:
                nom_c = fila_pers.iloc[0]["nom_str"]
                # --- AQUÍ ESTÁ LA SOLUCIÓN ---
                ape_c = str(fila_pers.iloc[0].get("apellidos", "")).strip()
                nom_p_c = str(fila_pers.iloc[0].get("nombres", "")).strip()
                # -----------------------------

               # --- NUEVA LÓGICA DE FOTO (Revisada para mayor tamaño y ajuste perfecto) ---
                # Buscamos la columna "foto" (minúscula o mayúscula)
                link_foto_raw = fila_pers.iloc[0].get("foto", fila_pers.iloc[0].get("FOTO", ""))
                
                # Transformamos el link (si usas Postimages/Blogger, la función lo dejará igual)
                if pd.notnull(link_foto_raw) and str(link_foto_raw).strip() != "":
                    foto_directa = obtener_link_directo_drive(str(link_foto_raw).strip())
                else:
                    foto_directa = None

                # Renderizamos la cabecera con FOTO MÁS GRANDE y AJUSTE PERFECTO
                if foto_directa:
                    st.markdown(f"""
                        <style>
                        /* Nueva clase para foto más grande que oculta los bordes guindos */
                        .foto-perfil-large {{
                            width: 110px;
                            height: 110px;
                            border-radius: 50%; 
                            object-fit: cover; 
                            object-position: center;
                            border: 4px solid #FFD700;
                            margin-right: 20px;
                            box-shadow: 0 4px 10px rgba(0,0,0,0.3);
                            transition: transform 0.2s ease-in-out;
                        }}
                        /* Efecto hover ligero, pero SIN clic y SIN abrir ventana */
                        .foto-perfil-large:hover {{
                            transform: scale(1.08);
                        }}
                        </style>
                        <div style='border-bottom: 2px solid #FFD700; padding-bottom: 15px; margin-bottom: 25px; display: flex; align-items: center;'>
                            <img src='{foto_directa}' class='foto-perfil-large' onerror="this.style.display='none'; document.getElementById('avatar-{dni_buscado}').style.display='block';">
                            <h1 id='avatar-{dni_buscado}' style='color: white; margin: 0; margin-right: 15px; font-size: 3em; display: none;'>👤</h1>
                            <h1 style='color: #FFD700; margin: 0; font-size: 2.5em;'>{nom_c}</h1>
                        </div>
                    """, unsafe_allow_html=True)
                else:
                    # Versión por defecto si no hay foto
                    st.markdown(f"""
                        <div style='border-bottom: 2px solid #FFD700; padding-bottom: 10px; margin-bottom: 20px; display: flex; align-items: center;'>
                            <h1 style='color: white; margin: 0; margin-right: 15px; font-size: 3em;'>👤</h1>
                            <h1 style='color: #FFD700; margin: 0; font-size: 2.5em;'>{nom_c}</h1>
                        </div>
                    """, unsafe_allow_html=True)
                                
                t_noms = ["Datos Generales", "Exp. Laboral", "Form. Académica", "Investigación", "Datos Familiares", "Contratos", "Vacaciones", "Otros Beneficios", "Méritos/Demer.", "Evaluación", "Liquidaciones"]
                h_keys = ["DATOS GENERALES", "EXP. LABORAL", "FORM. ACADEMICA", "INVESTIGACION", "DATOS FAMILIARES", "CONTRATOS", "VACACIONES", "OTROS BENEFICIOS", "MERITOS Y DEMERITOS", "EVALUACION DEL DESEMPEÑO", "LIQUIDACIONES"]

                tabs = st.tabs(t_noms)

                for i, tab in enumerate(tabs):
                    h_name = h_keys[i]
                    with tab:
                        if h_name in dfs and "dni" in dfs[h_name].columns:
                            c_df = dfs[h_name][dfs[h_name]["dni"] == dni_buscado]
                        else:
                            c_df = pd.DataFrame(columns=COLUMNAS.get(h_name, []))

                        if h_name == "CONTRATOS":
                            df_contratos = dfs["CONTRATOS"][dfs["CONTRATOS"]["dni"] == dni_buscado]
                            if not df_contratos.empty:
                                st.markdown("""
                                    <style>
                                    [data-testid="stDownloadButton"] button { background-color: #FFD700 !important; border: 2px solid #4A0000 !important; }
                                    [data-testid="stDownloadButton"] button p { color: #4A0000 !important; font-weight: bold !important; font-size: 16px !important; }
                                    [data-testid="stDownloadButton"] button:hover { background-color: #ffffff !important; border: 2px solid #FFD700 !important; }
                                    </style>
                                """, unsafe_allow_html=True)
                                word_file = gen_word(nom_c, dni_buscado, df_contratos)
                                st.download_button("📄 Generar Certificado de Trabajo", data=word_file, file_name=f"Certificado_{dni_buscado}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                                st.markdown("<br>", unsafe_allow_html=True)

                        if h_name == "VACACIONES":
                            df_tc = df_contratos[df_contratos["tipo contrato"].astype(str).str.lower().str.contains("planilla", na=False)] if "df_contratos" in locals() else pd.DataFrame()
                            
                            detalles = []
                            dias_generados_totales = 0
                            dias_gozados_totales = pd.to_numeric(c_df["dias gozados"], errors='coerce').sum()

                            if not df_tc.empty:
                                df_tc_calc = df_tc.copy()
                                df_tc_calc['f_inicio_dt'] = pd.to_datetime(df_tc_calc['f_inicio'], errors='coerce')
                                df_tc_calc['f_fin_dt'] = pd.to_datetime(df_tc_calc['f_fin'], errors='coerce')
                                
                                start_global = df_tc_calc['f_inicio_dt'].min()
                                
                                if pd.notnull(start_global):
                                    start_global = start_global.date()
                                    curr_start = start_global
                                    
                                    while curr_start <= date.today():
                                        curr_end = (pd.to_datetime(curr_start) + pd.DateOffset(years=1) - pd.Timedelta(days=1)).date()
                                        days_in_p = 0
                                        
                                        for _, r in df_tc_calc.iterrows():
                                            c_start = r['f_inicio_dt'].date() if pd.notnull(r['f_inicio_dt']) else None
                                            c_end = r['f_fin_dt'].date() if pd.notnull(r['f_fin_dt']) else None
                                            if c_start and c_end:
                                                o_start = max(curr_start, c_start)
                                                o_end = min(curr_end, c_end, date.today())
                                                if o_start <= o_end: 
                                                    days_in_p += (o_end - o_start).days + 1
                                                
                                        # --- SOLUCIÓN: CÁLCULO PROPORCIONAL EXACTO ---
                                        # Obtenemos los días totales reales que tiene ese periodo (365 o 366 si cruza un bisiesto)
                                        total_dias_periodo = (curr_end - curr_start).days + 1
                                        
                                        # Nueva fórmula: garantizamos un máximo exacto de 30 días por año completo
                                        gen_p = round((days_in_p / total_dias_periodo) * 30, 2)
                                        # ---------------------------------------------
                                        
                                        p_name = f"{curr_start.year}-{curr_start.year+1}"
                                        
                                        goz_df = c_df[c_df["periodo"].astype(str).str.strip() == p_name]
                                        goz_p = pd.to_numeric(goz_df["dias gozados"], errors='coerce').sum()
                                        
                                        if gen_p > 0 or goz_p > 0:
                                            detalles.append({"Periodo": p_name, "Del": curr_start.strftime("%d/%m/%Y"), "Al": curr_end.strftime("%d/%m/%Y"), "Días Generados": gen_p, "Dias Gozados": goz_p, "Saldo": round(gen_p - goz_p, 2)})
                                        
                                        dias_generados_totales += gen_p
                                        curr_start = (pd.to_datetime(curr_start) + pd.DateOffset(years=1)).date()

                            saldo_v = round(dias_generados_totales - dias_gozados_totales, 2)

                            st.markdown(f"""
                            <div style="display: flex; gap: 15px; margin-bottom: 20px;">
                                <div style="flex: 1; background-color: #4A0000; padding: 20px; border-radius: 10px; text-align: center; border: 2px solid #FFD700;"><h2 style="color: #FFD700; margin: 0; font-size: 2.5em;">{dias_generados_totales:.2f}</h2><p style="color: #FFFFFF; margin: 0; font-weight: bold; font-size: 1.1em;">Días Generados Totales</p></div>
                                <div style="flex: 1; background-color: #4A0000; padding: 20px; border-radius: 10px; text-align: center; border: 2px solid #FFD700;"><h2 style="color: #FFD700; margin: 0; font-size: 2.5em;">{dias_gozados_totales:.2f}</h2><p style="color: #FFFFFF; margin: 0; font-weight: bold; font-size: 1.1em;">Dias Gozados</p></div>
                                <div style="flex: 1; background-color: #4A0000; padding: 20px; border-radius: 10px; text-align: center; border: 2px solid #FFD700;"><h2 style="color: #FFD700; margin: 0; font-size: 2.5em;">{saldo_v:.2f}</h2><p style="color: #FFFFFF; margin: 0; font-weight: bold; font-size: 1.1em;">Saldo Disponible</p></div>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            if detalles:
                                st.markdown("<h4 style='color: #FFD700;'>Desglose por Periodos</h4>", unsafe_allow_html=True)
                                div_table = "<div style='display: flex; flex-direction: column; width: 100%; border: 2px solid #FFD700; border-radius: 8px; overflow: hidden; margin-bottom: 20px;'><div style='display: flex; background-color: #4A0000; color: #FFD700; font-weight: bold;'><div style='flex: 1; padding: 12px; text-align: center; border-right: 1px solid #FFD700;'>PERIODO</div><div style='flex: 1; padding: 12px; text-align: center; border-right: 1px solid #FFD700;'>DEL</div><div style='flex: 1; padding: 12px; text-align: center; border-right: 1px solid #FFD700;'>AL</div><div style='flex: 1; padding: 12px; text-align: center; border-right: 1px solid #FFD700;'>DÍAS GENERADOS</div><div style='flex: 1; padding: 12px; text-align: center; border-right: 1px solid #FFD700;'>DIAS GOZADOS</div><div style='flex: 1; padding: 12px; text-align: center;'>SALDO</div></div>"
                                for d in detalles:
                                    div_table += f"<div style='display: flex; background-color: #FFF9C4; color: #4A0000; font-weight: bold; border-top: 1px solid #FFD700;'><div style='flex: 1; padding: 10px; text-align: center; border-right: 1px solid #FFD700;'>{d['Periodo']}</div><div style='flex: 1; padding: 10px; text-align: center; border-right: 1px solid #FFD700;'>{d['Del']}</div><div style='flex: 1; padding: 10px; text-align: center; border-right: 1px solid #FFD700;'>{d['Al']}</div><div style='flex: 1; padding: 10px; text-align: center; border-right: 1px solid #FFD700;'>{d['Días Generados']:.2f}</div><div style='flex: 1; padding: 10px; text-align: center; border-right: 1px solid #FFD700;'>{d['Dias Gozados']:.2f}</div><div style='flex: 1; padding: 10px; text-align: center;'>{d['Saldo']:.2f}</div></div>"
                                div_table += "</div>"
                                st.markdown(div_table, unsafe_allow_html=True)

                        vst = c_df.copy()
                        
                        cols_ocultar = [c for c in vst.columns if c.lower() in ["apellidos y nombres", "apellidos", "nombres"]]
                        vst = vst.drop(columns=cols_ocultar)

                        col_conf = {}
                        for col in vst.columns:
                            if "fecha" in col.lower() or "f_" in col.lower():
                                vst[col] = pd.to_datetime(vst[col], errors='coerce').dt.date
                                # Volvemos a usar date() ahora que ya no hay reglas duplicadas
                                col_conf[str(col).upper()] = st.column_config.DateColumn(
                                    format="DD/MM/YYYY",
                                    min_value=date(1950, 1, 1),
                                    max_value=date(2100, 12, 31)
                                )
                            elif col.lower().strip() == "periodo":
                                vst[col] = vst[col].astype(str)
                                col_conf[str(col).upper()] = st.column_config.TextColumn()

                        vst.columns = [str(col).upper() for col in vst.columns]
                            
                        # Eliminamos duplicados de la base
                        vst = vst.loc[:, ~vst.columns.duplicated()]
            
                      # =========================================================
                        # 1. SI ES DATOS GENERALES -> DISEÑO TIPO FICHA (TARJETA)
                        # =========================================================
                        if h_name == "DATOS GENERALES" and not vst.empty:
                            ficha = vst.iloc[0]
                            
                            def get_val(names):
                                # Limpieza interna para asegurar que encuentre las columnas
                                ficha_clean = {str(k).lower().replace('á','a').replace('é','e').replace('í','i').replace('ó','o').replace('ú','u').replace('_',' '): v 
                                               for k, v in ficha.to_dict().items()}
                                for name in names:
                                    clean_name = name.lower().replace('á','a').replace('é','e').replace('í','i').replace('ó','o').replace('ú','u').replace('_',' ')
                                    val = ficha_clean.get(clean_name)
                                    if pd.notnull(val) and str(val).strip() not in ["", "-", "0", "nan"]: 
                                        return str(val)
                                return "-"

                            # Asignación de variables desde tu Google Sheets
                            sede = get_val(['SEDE'])
                            sexo = get_val(['SEXO'])
                            est_civil = get_val(['ESTADO CIVIL', 'ESTADO_CIVIL'])
                            f_nac = get_val(['FECHA DE NACIMIENTO', 'NACIMIENTO'])
                            edad = get_val(['EDAD'])
                            telefono = get_val(['CELULAR', 'TELEFONO', 'TELÉFONO'])
                            correo = get_val(['CORREO', 'EMAIL', 'CORREO ELECTRONICO'])
                            direccion = get_val(['DIRECCION', 'DIRECCIÓN', 'DOMICILIO'])
                            
                            dir_display = "-"
                            if direccion != "-":
                                query_map = direccion.replace(" ", "+")
                                link_mapa = f"https://www.google.com/maps/search/?api=1&query={query_map}"
                                dir_display = f'<a href="{link_mapa}" target="_blank" style="color: #4da3ff; text-decoration: none; font-weight: bold;">📍 {direccion} (Ver en Google Maps 🗺️)</a>'

                            st.markdown(f"""
                            <div style="background-color: rgba(255, 215, 0, 0.05); padding: 25px; border-radius: 15px; border: 2px solid #FFD700; color: inherit; font-family: sans-serif;">
                                <h2 style="margin-top:0; color: #FFD700; border-bottom: 1px solid rgba(255,215,0,0.3); padding-bottom:10px;">🪪 Expediente del Personal</h2>
                                <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); gap: 20px; margin-top: 15px;">
                                    <div><p style="margin:0; font-size: 0.85em; opacity: 0.7;">📍 SEDE</p><p style="margin:0; font-weight: bold; font-size: 1.1em;">{sede}</p></div>
                                    <div><p style="margin:0; font-size: 0.85em; opacity: 0.7;">🚻 SEXO</p><p style="margin:0; font-weight: bold; font-size: 1.1em;">{sexo}</p></div>
                                    <div><p style="margin:0; font-size: 0.85em; opacity: 0.7;">💍 ESTADO CIVIL</p><p style="margin:0; font-weight: bold; font-size: 1.1em;">{est_civil}</p></div>
                                    <div><p style="margin:0; font-size: 0.85em; opacity: 0.7;">🎂 F. NACIMIENTO</p><p style="margin:0; font-weight: bold; font-size: 1.1em;">{f_nac}</p></div>
                                    <div><p style="margin:0; font-size: 0.85em; opacity: 0.7;">🔢 EDAD ACTUAL</p><p style="margin:0; font-weight: bold; font-size: 1.1em;">{edad} años</p></div>
                                    <div><p style="margin:0; font-size: 0.85em; opacity: 0.7;">📱 TELÉFONO / CELULAR</p><p style="margin:0; font-weight: bold; font-size: 1.1em;">{telefono}</p></div>
                                </div>
                                <div style="margin-top: 25px; padding-top: 15px; border-top: 1px dashed rgba(255,215,0,0.3);">
                                    <div style="margin-bottom: 15px;">
                                        <p style="margin:0; font-size: 0.85em; opacity: 0.7;">📧 CORREO ELECTRÓNICO</p>
                                        <p style="margin:0; font-weight: bold; font-size: 1.1em;">{correo}</p>
                                    </div>
                                    <div>
                                        <p style="margin:0; font-size: 0.85em; opacity: 0.7;">🏠 DIRECCIÓN DE DOMICILIO</p>
                                        <p style="margin:0; font-size: 1.1em;">{dir_display}</p>
                                    </div>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            st.write("")
                            sel = vst.head(1)
                                                                                                   
                        # =========================================================
                        # 2. SI ES CUALQUIER OTRA PESTAÑA -> DISEÑO DE TABLA NORMAL
                        # =========================================================
                        else:
                            if "SEL" not in vst.columns:
                                vst.insert(0, "SEL", False)
                                
                            # Dejamos un solo "DIAS GENERADOS" sin tilde
                            columnas_basura = ["DNI", "FECHA DE INICIO", "FECHA DE FIN", "DIAS GENERADOS", "SALDO"]
                            for col in columnas_basura:
                                if col in vst.columns:
                                    col_conf[col] = None
                                    
                            # Dejamos un solo "DIAS GOZADOS" sin tilde
                            cols_importantes = ["SEL", "PERIODO", "F_INICIO", "F_FIN", "DIAS GOZADOS"]
                            cols_finales = [c for c in cols_importantes if c in vst.columns] + [c for c in vst.columns if c not in cols_importantes]
                                
                            # TRUCO ANTIFALLOS: Eliminamos cualquier duplicado accidental en la lista final
                            cols_finales = list(dict.fromkeys(cols_finales))
                            vst = vst[cols_finales]

                            # ==========================================
                            # NUEVO DISEÑO: INVESTIGACIÓN (TARJETAS VISUALES)
                            # ==========================================
                            if h_name == "INVESTIGACION":
                                col_izq, col_der = st.columns([2, 1])
                                
                                # --- LÓGICA DE LECTURA ---
                                df_inv = dfs.get("INVESTIGACION", pd.DataFrame())
                                col_dni_inv = "dni" if "dni" in df_inv.columns else "DNI"
                                
                                inv_empleado = pd.DataFrame()
                                if not df_inv.empty and col_dni_inv in df_inv.columns:
                                    inv_empleado = df_inv[df_inv[col_dni_inv].astype(str) == str(dni_buscado)]
                                
                                conteo_pub = conteo_fondos = conteo_sem = 0
                                
                                # --- COLUMNA IZQUIERDA: TARJETAS ---
                                with col_izq:
                                    st.markdown("<h3 style='color: #FFD700;'>🔬 Registro de Actividades de Investigación</h3>", unsafe_allow_html=True)
                                    
                                    if inv_empleado.empty:
                                        st.markdown("<p style='color:#DDDDDD;'>No hay registros de investigación para este colaborador.</p>", unsafe_allow_html=True)
                                    else:
                                        for idx, row in inv_empleado.iterrows():
                                            tipo = str(row.get('tipo de registro', row.get('TIPO DE REGISTRO', 'Otro'))).strip()
                                            
                                            # Tarjetas según el tipo
                                            if tipo == "Datos Generales (CTI Vitae / RENACYT)":
                                                renacyt = row.get('codigo renacyt', 'N/A')
                                                nivel = row.get('nivel renacyt', 'N/A')
                                                link = row.get('enlace cti vitae', '#')
                                                
                                                st.markdown(f"""
                                                <div style='background-color: #E8F4F8; padding: 15px; border-radius: 8px; border-left: 6px solid #00AEEF; margin-bottom: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); border: 1px solid #CCCCCC;'>
                                                    <div style='color: #000000; font-size: 1.1em; font-weight: bold; margin-bottom: 5px;'>👤 Perfil CTI Vitae / RENACYT</div>
                                                    <div style='color: #222222; font-size: 0.95em;'>
                                                        <strong>Código:</strong> {renacyt} <br>
                                                        <strong>Nivel RENACYT:</strong> {nivel} <br>
                                                        <a href="{link}" target="_blank" style="color: #00AEEF; text-decoration: none;">🔗 Ver Perfil CTI Vitae</a>
                                                    </div>
                                                </div>
                                                """, unsafe_allow_html=True)
                                                
                                            elif tipo == "Publicación Científica":
                                                conteo_pub += 1
                                                titulo = row.get('titulo de publicacion', 'N/A')
                                                bd = row.get('base de datos', 'N/A')
                                                revista = row.get('nombre de revista', 'N/A')
                                                anio = row.get('año de publicacion', 'N/A')
                                                
                                                st.markdown(f"""
                                                <div style='background-color: #F9F6EE; padding: 15px; border-radius: 8px; border-left: 6px solid #FF8C00; margin-bottom: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); border: 1px solid #CCCCCC;'>
                                                    <div style='color: #000000; font-size: 1.1em; font-weight: bold; margin-bottom: 5px;'>📄 Publicación: {titulo}</div>
                                                    <div style='color: #222222; font-size: 0.95em;'>
                                                        <strong>Revista:</strong> {revista} ({anio}) <br>
                                                        <strong>Indexación:</strong> {bd} - <strong>Cuartil:</strong> {row.get('cuartil', 'N/A')}
                                                    </div>
                                                </div>
                                                """, unsafe_allow_html=True)
                                                
                                            elif tipo == "Fondo Concursable":
                                                conteo_fondos += 1
                                                titulo_proy = row.get('nombre del proyecto', 'N/A')
                                                entidad = row.get('entidad financiadora', 'N/A')
                                                estado = row.get('estado del proyecto', 'N/A')
                                                
                                                st.markdown(f"""
                                                <div style='background-color: #F4FDE8; padding: 15px; border-radius: 8px; border-left: 6px solid #4CAF50; margin-bottom: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); border: 1px solid #CCCCCC;'>
                                                    <div style='color: #000000; font-size: 1.1em; font-weight: bold; margin-bottom: 5px;'>💰 Proyecto Financiado: {titulo_proy}</div>
                                                    <div style='color: #222222; font-size: 0.95em;'>
                                                        <strong>Entidad:</strong> {entidad} <br>
                                                        <strong>Rol:</strong> {row.get('rol en el proyecto', 'N/A')} <br>
                                                        <strong>Estado:</strong> <span style="color: {'green' if estado=='Finalizado' else 'blue'}; font-weight: bold;">{estado}</span>
                                                    </div>
                                                </div>
                                                """, unsafe_allow_html=True)
                                                
                                            elif tipo == "Semillero de Investigación":
                                                conteo_sem += 1
                                                nombre_sem = row.get('nombre del semillero', 'N/A')
                                                rol_sem = row.get('rol en el semillero', 'N/A')
                                                estado_sem = row.get('estado del semillero', 'N/A')
                                                
                                                st.markdown(f"""
                                                <div style='background-color: #F8E8F8; padding: 15px; border-radius: 8px; border-left: 6px solid #9C27B0; margin-bottom: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); border: 1px solid #CCCCCC;'>
                                                    <div style='color: #000000; font-size: 1.1em; font-weight: bold; margin-bottom: 5px;'>🌱 Semillero: {nombre_sem}</div>
                                                    <div style='color: #222222; font-size: 0.95em;'>
                                                        <strong>Rol:</strong> {rol_sem} <br>
                                                        <strong>Estado:</strong> {estado_sem}
                                                    </div>
                                                </div>
                                                """, unsafe_allow_html=True)

                                # --- COLUMNA DERECHA: DASHBOARD DE RESUMEN ---
                                with col_der:
                                    st.markdown("<h3 style='color: #FFD700;'>📊 Resumen</h3>", unsafe_allow_html=True)
                                    html_resumen_inv = f"""
                                    <div style='background-color: #4A0000; padding: 20px; border-radius: 10px; border: 2px solid #FFD700; box-shadow: 2px 2px 10px rgba(0,0,0,0.5); position: sticky; top: 50px;'>
                                        <h4 style='color: #FFD700; margin-bottom: 15px; text-align: center; border-bottom: 1px solid #FFD700; padding-bottom: 10px;'>Impacto Científico</h4>
                                        <div style='margin-bottom: 15px;'>
                                            <p style='margin: 0; color: #FFFFFF; font-size: 0.9em;'>📄 Publicaciones Indexadas</p>
                                            <p style='margin: 0; color: #FF8C00; font-size: 1.4em; font-weight: bold;'>{conteo_pub}</p>
                                        </div>
                                        <div style='margin-bottom: 15px;'>
                                            <p style='margin: 0; color: #FFFFFF; font-size: 0.9em;'>💰 Fondos Concursables</p>
                                            <p style='margin: 0; color: #4CAF50; font-size: 1.4em; font-weight: bold;'>{conteo_fondos}</p>
                                        </div>
                                        <div style='margin-bottom: 15px;'>
                                            <p style='margin: 0; color: #FFFFFF; font-size: 0.9em;'>🌱 Semilleros Liderados</p>
                                            <p style='margin: 0; color: #9C27B0; font-size: 1.4em; font-weight: bold;'>{conteo_sem}</p>
                                        </div>
                                    </div>
                                    """
                                    st.markdown(html_resumen_inv, unsafe_allow_html=True)
                                
                                # TABLA DE EDICIÓN OCULTA (Para que siga funcionando tu sistema de SEL)
                                st.markdown("<br>", unsafe_allow_html=True)
                                with st.expander("⚙️ Clic aquí para Editar o Eliminar un Registro de Investigación"):
                                    st.markdown("<p style='color:#DDDDDD;'>Activa la casilla <b>SEL</b> para modificar o eliminar un registro.</p>", unsafe_allow_html=True)
                                    st.markdown("""<style>[data-testid="stDataEditor"] { border: 2px solid #FFD700 !important; border-radius: 10px !important; }</style>""", unsafe_allow_html=True)
                                    ed = st.data_editor(vst, hide_index=True, use_container_width=True, column_config=col_conf, key=f"ed_{h_name}_oculta")
                                    sel = ed[ed["SEL"] == True]

                           # ==========================================
                            # NUEVO DISEÑO: EXPERIENCIA LABORAL Y CÁLCULOS
                            # ==========================================
                            if h_name == "EXP. LABORAL":
                                # Dividimos la pantalla: 2/3 para tarjetas, 1/3 para el resumen
                                col_izq, col_der = st.columns([2, 1])
                                
                                df_contratos = dfs.get("CONTRATOS", pd.DataFrame())
                                col_dni_contratos = "DNI" if "DNI" in df_contratos.columns else "dni"
                                
                                contratos_empleado = pd.DataFrame()
                                if not df_contratos.empty and col_dni_contratos in df_contratos.columns:
                                    contratos_empleado = df_contratos[df_contratos[col_dni_contratos] == str(dni_buscado)]
                                
                                # --- LÓGICA DE CÁLCULO DE TIEMPO Y FORMATO DE FECHAS ---
                                meses_docente = 0
                                meses_admin = 0
                                
                                def calcular_meses(f_ini, f_fin):
                                    try:
                                        inicio = pd.to_datetime(f_ini, errors='coerce')
                                        fin = pd.to_datetime(f_fin, errors='coerce')
                                        if pd.isna(inicio) or pd.isna(fin): return 0
                                        return max(0, int((fin - inicio).days / 30.44))
                                    except:
                                        return 0
                                        
                                def dar_formato_fecha(fecha_str):
                                    """Convierte cualquier fecha al formato bonito DD/MM/YYYY"""
                                    try:
                                        if pd.isna(fecha_str) or str(fecha_str).strip() == "" or str(fecha_str) == "NaT": 
                                            return "N/A"
                                        return pd.to_datetime(fecha_str).strftime('%d/%m/%Y')
                                    except:
                                        return str(fecha_str)

                                # ---------------------------------------
                                # COLUMNA IZQUIERDA: TARJETAS VISUALES
                                # ---------------------------------------
                                with col_izq:
                                    st.markdown("<h3 style='color: #FFD700;'>🏢 Experiencia Interna (Universidad Roosevelt)</h3>", unsafe_allow_html=True)
                                    if contratos_empleado.empty:
                                        st.markdown("<p style='color:#DDDDDD;'>No hay contratos internos registrados.</p>", unsafe_allow_html=True)
                                    else:
                                        for idx, row in contratos_empleado.iterrows():
                                            f_ini = row.get('f_inicio', row.get('F_INICIO', 'N/A'))
                                            f_fin = row.get('f_fin', row.get('F_FIN', 'N/A'))
                                            
                                            # Aplicamos formato bonito
                                            f_ini_str = dar_formato_fecha(f_ini)
                                            f_fin_str = dar_formato_fecha(f_fin)
                                            
                                            puesto = row.get('cargo', row.get('CARGO', row.get('PUESTO', 'N/A')))
                                            tipo_trabajador_raw = str(row.get('TIPO DE TRABAJADOR', row.get('tipo de trabajador', 'Administrativo')))
                                            tipo_exp = "Docente" if "docente" in tipo_trabajador_raw.lower() else "Administrativo"
                                            
                                            meses_calc = calcular_meses(f_ini, f_fin)
                                            if tipo_exp == "Docente": meses_docente += meses_calc
                                            else: meses_admin += meses_calc

                                            # Fondo Blanco Hueso (#F9F6EE)
                                            st.markdown(f"""
                                            <div style='background-color: #F9F6EE; padding: 15px; border-radius: 8px; border-left: 6px solid #4A0000; margin-bottom: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); border: 1px solid #CCCCCC;'>
                                                <div style='color: #000000; font-size: 1.1em; font-weight: bold; margin-bottom: 5px;'>{puesto} <span style='font-size: 0.85em; color: #555555;'>(Interno - {tipo_exp})</span></div>
                                                <div style='color: #222222; font-size: 0.95em;'>
                                                    <strong>Lugar:</strong> Universidad Roosevelt <br>
                                                    <strong>Periodo:</strong> {f_ini_str} hasta {f_fin_str} <br>
                                                    <strong>Tipo de Contrato:</strong> {row.get('tipo contrato', row.get('TIPO CONTRATO', 'N/A'))}
                                                </div>
                                            </div>
                                            """, unsafe_allow_html=True)
                                    
                                    st.markdown("<h3 style='color: #FFD700; margin-top: 20px;'>💼 Experiencia Externa Registrada</h3>", unsafe_allow_html=True)
                                    if vst.empty:
                                        st.markdown("<p style='color:#DDDDDD;'>No hay experiencia externa registrada.</p>", unsafe_allow_html=True)
                                    else:
                                        for idx, row in vst.iterrows():
                                            f_ini = row.get('FECHA DE INICIO', row.get('fecha de inicio', 'N/A'))
                                            f_fin = row.get('FECHA DE FIN', row.get('fecha de fin', 'N/A'))
                                            
                                            # Aplicamos formato bonito
                                            f_ini_str = dar_formato_fecha(f_ini)
                                            f_fin_str = dar_formato_fecha(f_fin)
                                            
                                            tipo_exp_raw = str(row.get('TIPO DE EXPERIENCIA', row.get('tipo de experiencia', 'Administrativo')))
                                            tipo_exp = "Docente" if "docente" in tipo_exp_raw.lower() else "Administrativo"
                                            
                                            meses_calc = calcular_meses(f_ini, f_fin)
                                            if tipo_exp == "Docente": meses_docente += meses_calc
                                            else: meses_admin += meses_calc

                                            puesto_ext = row.get('PUESTO', row.get('puesto', 'N/A'))
                                            lugar_ext = row.get('LUGAR', row.get('lugar', 'N/A'))
                                            motivo_ext = row.get('MOTIVO DE CESE', row.get('motivo de cese', 'N/A'))

                                            # Fondo Blanco Hueso (#F9F6EE)
                                            st.markdown(f"""
                                            <div style='background-color: #F9F6EE; padding: 15px; border-radius: 8px; border-left: 6px solid #004A80; margin-bottom: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); border: 1px solid #CCCCCC;'>
                                                <div style='color: #000000; font-size: 1.1em; font-weight: bold; margin-bottom: 5px;'>{puesto_ext} <span style='font-size: 0.85em; color: #555555;'>({tipo_exp.capitalize()})</span></div>
                                                <div style='color: #222222; font-size: 0.95em;'>
                                                    <strong>Lugar:</strong> {lugar_ext} <br>
                                                    <strong>Periodo:</strong> {f_ini_str} hasta {f_fin_str} <br>
                                                    <strong>Motivo de cese:</strong> {motivo_ext}
                                                </div>
                                            </div>
                                            """, unsafe_allow_html=True)
                                            
                                # ---------------------------------------
                                # COLUMNA DERECHA: DASHBOARD DE TIEMPO
                                # ---------------------------------------
                                with col_der:
                                    def formato_tiempo(total_meses):
                                        anios = total_meses // 12
                                        meses = total_meses % 12
                                        if anios > 0 and meses > 0: return f"{anios} años y {meses} meses"
                                        elif anios > 0: return f"{anios} años"
                                        elif meses > 0: return f"{meses} meses"
                                        else: return "0 meses"

                                    st.markdown("<h3 style='color: #FFD700;'>📊 Resumen</h3>", unsafe_allow_html=True)
                                    
                                    html_resumen = f"""
                                    <div style='background-color: #4A0000; padding: 20px; border-radius: 10px; border: 2px solid #FFD700; box-shadow: 2px 2px 10px rgba(0,0,0,0.5); position: sticky; top: 50px;'>
                                        <h4 style='color: #FFD700; margin-bottom: 15px; text-align: center; border-bottom: 1px solid #FFD700; padding-bottom: 10px;'>Tiempo Total Calculado</h4>
                                        <div style='margin-bottom: 15px;'>
                                            <p style='margin: 0; color: #FFFFFF; font-size: 0.9em;'>👨‍🏫 Como Docente</p>
                                            <p style='margin: 0; color: #FFD700; font-size: 1.2em; font-weight: bold;'>{formato_tiempo(meses_docente)}</p>
                                        </div>
                                        <div style='margin-bottom: 15px;'>
                                            <p style='margin: 0; color: #FFFFFF; font-size: 0.9em;'>💼 Como Administrativo</p>
                                            <p style='margin: 0; color: #FFD700; font-size: 1.2em; font-weight: bold;'>{formato_tiempo(meses_admin)}</p>
                                        </div>
                                        <div style='margin-top: 15px; padding-top: 10px; border-top: 1px solid #FFD700;'>
                                            <p style='margin: 0; color: #FFFFFF; font-size: 0.9em;'>🌟 Experiencia General</p>
                                            <p style='margin: 0; color: #00FF00; font-size: 1.4em; font-weight: bold;'>{formato_tiempo(meses_docente + meses_admin)}</p>
                                        </div>
                                    </div>
                                    """
                                    st.markdown(html_resumen, unsafe_allow_html=True)

                                # ==========================================
                                    # MÓDULO INTELIGENTE: ANÁLISIS DE PERFIL Y LÍNEA DE CARRERA
                                    # ==========================================
                                    st.markdown("<br>", unsafe_allow_html=True)
                                    st.markdown("<h4 style='color: #4A0000;'>🎯 Plan de Carrera (Análisis SUNEDU/Estatuto)</h4>", unsafe_allow_html=True)

                                    # 1. RECOLECCIÓN DE DATOS DEL TRABAJADOR
                                    es_doctor = False
                                    es_maestro = False
                                    tiene_renacyt = False
                                    total_publicaciones = 0
                                    anios_docencia = meses_docente // 12
                                    
                                    # Analizamos Formación Académica
                                    df_acad = dfs.get("FORM. ACADEMICA", pd.DataFrame())
                                    if not df_acad.empty and "dni" in df_acad.columns:
                                        acad_emp = df_acad[df_acad["dni"].astype(str) == str(dni_buscado)]
                                        for idx, row in acad_emp.iterrows():
                                            grado = str(row.get('grado o titulo obtenido', '')).upper()
                                            if "DOCTOR" in grado: es_doctor = True
                                            if "MAGISTER" in grado or "MAESTRO" in grado or "MAESTRIA" in grado: es_maestro = True

                                    # Analizamos Investigación
                                    df_inv = dfs.get("INVESTIGACION", pd.DataFrame())
                                    if not df_inv.empty and "dni" in df_inv.columns:
                                        inv_emp = df_inv[df_inv["dni"].astype(str) == str(dni_buscado)]
                                        for idx, row in inv_emp.iterrows():
                                            tipo = str(row.get('tipo de registro', ''))
                                            nivel_renacyt = str(row.get('nivel renacyt', 'No tiene'))
                                            
                                            if "Datos Generales" in tipo and nivel_renacyt != "No tiene":
                                                tiene_renacyt = True
                                            if "Publicación Científica" in tipo:
                                                total_publicaciones += 1

                                    # 2. SISTEMA DE PUNTUACIÓN AUTOMÁTICA (Baremo UPHFR)
                                    puntos_formacion = 40 if es_doctor else (25 if es_maestro else 10)
                                    puntos_investigacion = (15 if tiene_renacyt else 0) + (total_publicaciones * 5)
                                    if puntos_investigacion > 30: puntos_investigacion = 30 # Tope máximo
                                    puntos_experiencia = anios_docencia * 2
                                    if puntos_experiencia > 30: puntos_experiencia = 30 # Tope máximo

                                    puntaje_total = puntos_formacion + puntos_investigacion + puntos_experiencia

                                    # Contenedor para el Plan de Carrera
                                    with st.container():
                                        st.markdown(f"""
                                        <div style='background-color: #F9F6EE; padding: 15px; border-radius: 8px; border: 1px solid #CCCCCC;'>
                                            <div style='color: #004A80; font-size: 1.1em; font-weight: bold; margin-bottom: 10px;'>📊 Nivel de Competitividad Institucional: {puntaje_total}/100 pts</div>
                                            <div style='font-size: 0.85em; color: #555555; margin-bottom: 5px;'>🎓 Formación: {puntos_formacion}/40 | 💼 Exp: {puntos_experiencia}/30 | 🔬 Inv: {puntos_investigacion}/30</div>
                                        </div>
                                        """, unsafe_allow_html=True)
                                        
                                        st.progress(puntaje_total / 100)
                                        st.markdown("<br>", unsafe_allow_html=True)

                                        # Acordeón para ver los Match de Cargos
                                        with st.expander("🏛️ Ver Idoneidad para Cargos UPHFR (Estatuto)"):
                                            # Evaluación para RECTOR (Art. 64)
                                            if es_doctor and anios_docencia >= 5: 
                                                st.success("✅ **Rector / Vicerrector:** CUMPLE (Tiene Grado de Doctor y experiencia requerida según Art. 64).")
                                            else:
                                                st.error("❌ **Rector / Vicerrector:** NO CUMPLE (Requiere Grado de Doctor y 5+ años de experiencia docente).")
                                                
                                            # Evaluación para DIRECTOR DE POSGRADO (Art. 25 y 28)
                                            if es_doctor:
                                                st.success("✅ **Director de Escuela de Posgrado:** CUMPLE (Tiene el grado máximo que otorga la unidad).")
                                            else:
                                                st.error("❌ **Director de Escuela de Posgrado:** NO CUMPLE (Requiere Grado de Doctor).")

                                            # Evaluación para DOCENTE INVESTIGADOR
                                            if tiene_renacyt:
                                                st.success("✅ **Docente Investigador:** CUMPLE (Cuenta con clasificación RENACYT activa).")
                                            else:
                                                st.warning("⚠️ **Docente Investigador:** EN PROCESO (Requiere obtener clasificación RENACYT).")

                                            # Evaluación para DOCENTE DE PREGRADO y Match de Carreras
                                            if es_maestro or es_doctor:
                                                st.success("✅ **Docente Universitario (Pregrado):** CUMPLE (Cuenta con Grado de Maestro o superior).")
                                                
                                                # --- MOTOR DE MATCH DE ESPECIALIDADES ---
                                                st.markdown("<p style='color: #004A80; font-weight: bold; margin-top: 10px; margin-bottom: 5px;'>🔍 Análisis de Especialidad Docente:</p>", unsafe_allow_html=True)
                                                
                                                # 1. Recopilar texto a analizar (Grados y Puestos)
                                                texto_perfil = ""
                                                if not df_acad.empty and "dni" in df_acad.columns:
                                                    for _, row in acad_emp.iterrows():
                                                        texto_perfil += " " + str(row.get('grado o titulo obtenido', '')).upper()
                                                        texto_perfil += " " + str(row.get('especialidad', '')).upper() # Si tienes esta columna
                                                        
                                                # Añadimos la experiencia externa al texto de análisis
                                                if 'vst' in locals() and not vst.empty:
                                                    for _, row in vst.iterrows():
                                                        texto_perfil += " " + str(row.get('PUESTO', row.get('puesto', ''))).upper()
                                                        texto_perfil += " " + str(row.get('LUGAR', row.get('lugar', ''))).upper()

                                                # 2. Diccionario de Carreras UPHFR
                                                diccionario_carreras = {
                                                    "Enfermería": ["ENFERMER", "CUIDADO", "CLINIC"],
                                                    "Medicina Humana": ["MEDICIN", "MEDICO", "CIRUJAN", "CLINIC", "HOSPITAL"],
                                                    "Obstetricia": ["OBSTETR", "MATRON", "GINECOLOG"],
                                                    "Farmacia y Bioquímica": ["FARMAC", "BIOQUIMIC", "LABORATORI", "QUIMIC"],
                                                    "Psicología": ["PSICOLOG", "SALUD MENTAL", "TERAP"],
                                                    "Estomatología": ["ESTOMATOLOG", "ODONTOLOG", "DENTIS"],
                                                    "Derecho": ["DERECHO", "ABOGAD", "LEGAL", "JURIDIC", "LEY", "MAGISTRAD", "JUEZ", "FISCAL"],
                                                    "Administración": ["ADMINISTRAC", "GERENCI", "NEGOCIO", "EMPRES", "CONTABILIDAD", "ECONOMI"]
                                                }

                                                # 3. Analizar coincidencias
                                                match_salud = []
                                                match_empresariales = []
                                                
                                                for carrera, palabras in diccionario_carreras.items():
                                                    if any(palabra in texto_perfil for palabra in palabras):
                                                        if carrera in ["Enfermería", "Medicina Humana", "Obstetricia", "Farmacia y Bioquímica", "Psicología", "Estomatología"]:
                                                            match_salud.append(carrera)
                                                        else:
                                                            match_empresariales.append(carrera)

                                                # 4. Mostrar Resultados
                                                if not match_salud and not match_empresariales:
                                                    st.info("ℹ️ Perfil multidisciplinario. Se requiere revisión manual para asignar cursos específicos.")
                                                else:
                                                    if match_salud:
                                                        st.markdown(f"**🏥 C. de la Salud:** Apto para dictar en **{', '.join(match_salud)}**.")
                                                    if match_empresariales:
                                                        st.markdown(f"**🏢 C. Empresariales:** Apto para dictar en **{', '.join(match_empresariales)}**.")
                                                        
                                            else:
                                                st.error("❌ **Docente Universitario (Pregrado):** NO CUMPLE (La Ley exige mínimo Grado de Maestro).")
                               # ---------------------------------------
                                # TABLA DESPLEGABLE PARA EDICIÓN
                                # ---------------------------------------
                                st.markdown("<br>", unsafe_allow_html=True)
                                with st.expander("⚙️ Clic aquí para Editar o Eliminar Experiencia Externa"):
                                    st.markdown("<p style='color:#DDDDDD;'>Activa la casilla <b>SEL</b> para modificar o eliminar un registro.</p>", unsafe_allow_html=True)
                                    st.markdown("""<style>[data-testid="stDataEditor"] { border: 2px solid #FFD700 !important; border-radius: 10px !important; }</style>""", unsafe_allow_html=True)
                                    
                                    # Usamos directamente el col_conf global que ya arreglamos arriba
                                    ed = st.data_editor(vst, hide_index=True, use_container_width=True, column_config=col_conf, key=f"ed_{h_name}_oculta")
                                    
                                    # Convertimos los calendarios de vuelta a texto antes de que pasen a SEL
                                    for col in ed.columns:
                                        if "fecha" in col.lower() or "f_" in col.lower():
                                            ed[col] = ed[col].astype(str)
                                            ed[col] = ed[col].replace(["NaT", "None"], "") # Limpiamos vacíos

                                    sel = ed[ed["SEL"] == True]
                            # ==========================================
                            # NUEVO DISEÑO: CONTRATOS
                            # ==========================================
                            elif h_name == "CONTRATOS":
                                if vst.empty:
                                    st.markdown("<p style='color:#DDDDDD;'>No hay contratos registrados para este colaborador.</p>", unsafe_allow_html=True)
                                else:
                                    # Generamos tarjetas estéticas para cada contrato (estilo claro)
                                    for _, row in vst.iterrows():
                                        f_inicio = row.get('F_INICIO', 'N/A')
                                        f_fin = row.get('F_FIN', 'N/A')
                                        cargo = row.get('CARGO', 'N/A')
                                        tipo = row.get('TIPO CONTRATO', 'N/A')
                                        estado = row.get('ESTADO', 'N/A')
                                        
                                        # Verde si es activo, Rojo si es cesado
                                        color_borde = "#4CAF50" if str(estado).strip().upper() == "ACTIVO" else "#F44336"
                                        
                                        st.markdown(f"""
                                        <div style='background-color: #FFFFFF; padding: 15px; border-radius: 8px; border-left: 6px solid {color_borde}; margin-bottom: 10px; border: 1px solid #CCCCCC;'>
                                            <div style='color: #000000; font-size: 1.1em; font-weight: bold; margin-bottom: 5px;'>{cargo}</div>
                                            <div style='color: #111111; font-size: 0.95em;'>
                                                <strong>📅 Periodo:</strong> {f_inicio} hasta {f_fin} <br>
                                                <strong>📝 Tipo de Contrato:</strong> {tipo} <br>
                                                <strong>📌 Estado:</strong> <span style='color: {color_borde}; font-weight: bold;'>{estado}</span>
                                            </div>
                                        </div>
                                        """, unsafe_allow_html=True)

                                # TABLA DESPLEGABLE PARA EDICIÓN
                                st.markdown("<br>", unsafe_allow_html=True)
                                with st.expander("⚙️ Clic aquí para Editar o Eliminar Contratos"):
                                    st.markdown("<p style='color:#DDDDDD;'>Activa la casilla <b>SEL</b> en la tabla de abajo para modificar o eliminar un registro.</p>", unsafe_allow_html=True)
                                    st.markdown("""<style>[data-testid="stDataEditor"] { border: 2px solid #FFD700 !important; border-radius: 8px !important; }</style>""", unsafe_allow_html=True)
                                    ed = st.data_editor(vst, hide_index=True, use_container_width=True, column_config=col_conf, key=f"ed_{h_name}_oculta")
                                    sel = ed[ed["SEL"] == True]

                            # ==========================================
                            # NUEVO DISEÑO: VACACIONES
                            # ==========================================
                            elif h_name == "VACACIONES":
                                # TABLA DESPLEGABLE PARA EDICIÓN (La parte visual ya se calculó más arriba)
                                st.markdown("<br>", unsafe_allow_html=True)
                                with st.expander("⚙️ Clic aquí para ver el Detalle Completo, Editar o Eliminar Vacaciones"):
                                    st.markdown("<p style='color:#000000; background-color:#FFD700; padding:5px; border-radius:5px;'><b>Detalle de registros:</b> Activa la casilla <b>SEL</b> para modificar o eliminar.</p>", unsafe_allow_html=True)
                                    st.markdown("""<style>[data-testid="stDataEditor"] { border: 2px solid #FFD700 !important; border-radius: 8px !important; }</style>""", unsafe_allow_html=True)
                                    ed = st.data_editor(vst, hide_index=True, use_container_width=True, column_config=col_conf, key=f"ed_{h_name}_oculta")
                                    sel = ed[ed["SEL"] == True]
                            # ==========================================
                            # NUEVO DISEÑO: FORMACIÓN ACADÉMICA
                            # ==========================================
                            elif h_name == "FORM. ACADEMICA":
                                st.markdown("<h3 style='color: #FFD700; margin-bottom: 20px;'>🎓 Resumen de Formación Académica</h3>", unsafe_allow_html=True)
                                
                                col_tipo = "TIPO DE ESTUDIO" if "TIPO DE ESTUDIO" in vst.columns else "tipo de estudio"
                                
                                if not vst.empty and col_tipo in vst.columns:
                                    # 1. Filtramos primero los estudios sin grado, terminados o inconclusos
                                    mask_estudios = vst[col_tipo].str.contains("Terminados|Inconclusos|Sin grado", case=False, na=False)
                                    df_estudios = vst[mask_estudios]
                                    
                                    # 2. Filtramos los Grados/Títulos, pero EXCLUYENDO a los de arriba (usamos el símbolo ~ para negar)
                                    mask_grados = vst[col_tipo].str.contains("Grado|Títul|Titul", case=False, na=False) & ~mask_estudios
                                    df_grados = vst[mask_grados]
                                    
                                    # 3. Los demás se quedan igual
                                    df_especi = vst[vst[col_tipo].str.contains("Especialización|Especializaciones", case=False, na=False)]
                                    df_diplo = vst[vst[col_tipo].str.contains("Diplomado", case=False, na=False)]
                                    df_cursos = vst[vst[col_tipo].str.contains("Curso", case=False, na=False)]
                                else:
                                    df_grados = df_estudios = df_especi = df_diplo = df_cursos = pd.DataFrame()

                                # Función mejorada para evitar celdas vacías (N/A)
                                def get_val(r, opciones):
                                    for op in opciones:
                                        if op in r:
                                            val = r[op]
                                            if pd.notna(val) and str(val).strip() != "":
                                                return str(val)
                                    return "N/A"

                                # --- 1. GRADOS Y TÍTULOS ---
                                st.markdown("<h4 style='color: #FFD700; font-weight: bold; border-bottom: 2px solid #FFD700; padding-bottom: 5px;'>📜 Grados y Títulos</h4>", unsafe_allow_html=True)
                                if df_grados.empty:
                                    st.markdown("<p style='color:#DDDDDD;'>No hay grados o títulos registrados.</p>", unsafe_allow_html=True)
                                else:
                                    for _, row in df_grados.iterrows():
                                        grado = get_val(row, ['grado o titulo obtenido', 'GRADO O TITULO OBTENIDO', 'grado o título obtenido'])
                                        inst = get_val(row, ['institucion educativa', 'INSTITUCION EDUCATIVA', 'institución educativa'])
                                        mencion = get_val(row, ['mencion (especialidad / carrera / etc)', 'MENCION (ESPECIALIDAD / CARRERA / ETC)', 'mención'])
                                        anio = get_val(row, ['AÑO', 'año'])
                                        
                                        st.markdown(f"""
                                        <div style='background-color: #FFFFFF; padding: 15px; border-radius: 8px; border: 1px solid #CCCCCC; border-left: 6px solid #FFC107; margin-bottom: 10px;'>
                                            <div style='margin-bottom: 5px; color: #000000; font-size: 1.1em; font-weight: bold;'>{grado}</div>
                                            <div style='margin: 2px 0; color: #000000;'><strong>Institución:</strong> {inst}</div>
                                            <div style='margin: 2px 0; color: #000000;'><strong>Mención:</strong> {mencion}</div>
                                            <div style='margin: 2px 0; color: #000000;'><strong>Año:</strong> {anio}</div>
                                        </div>
                                        """, unsafe_allow_html=True)

                                col_izq_acad, col_der_acad = st.columns(2)

                                with col_izq_acad:
                                    # --- 2. ESTUDIOS TERMINADOS O INCONCLUSOS ---
                                    st.markdown("<h4 style='color: #FFD700; font-weight: bold; margin-top: 15px; border-bottom: 2px solid #FFD700; padding-bottom: 5px;'>🚧 Estudios Sin Grado / Inconclusos</h4>", unsafe_allow_html=True)
                                    if df_estudios.empty:
                                        st.markdown("<p style='color:#DDDDDD;'>No registrados.</p>", unsafe_allow_html=True)
                                    else:
                                        for _, row in df_estudios.iterrows():
                                            inst = get_val(row, ['institucion educativa', 'INSTITUCION EDUCATIVA', 'institución educativa'])
                                            mencion = get_val(row, ['mencion (especialidad / carrera / etc)', 'MENCION (ESPECIALIDAD / CARRERA / ETC)', 'mención'])
                                            anio = get_val(row, ['AÑO', 'año'])
                                            estado = get_val(row, ['ESTADO', 'estado'])
                                            
                                            st.markdown(f"""
                                            <div style='background-color: #FFFFFF; padding: 15px; border-radius: 8px; border: 1px solid #CCCCCC; border-left: 6px solid #FF5722; margin-bottom: 10px;'>
                                                <div style='margin-bottom: 5px; color: #000000; font-size: 1em; font-weight: bold;'>{mencion}</div>
                                                <div style='margin: 2px 0; color: #000000;'><strong>Institución:</strong> {inst}</div>
                                                <div style='margin: 2px 0; color: #000000;'><strong>Estado:</strong> <span style='color: #D84315; font-weight: bold;'>{estado}</span> | <strong>Año:</strong> {anio}</div>
                                            </div>
                                            """, unsafe_allow_html=True)

                                    # --- 3. ESPECIALIZACIONES ---
                                    st.markdown("<h4 style='color: #FFD700; font-weight: bold; margin-top: 15px; border-bottom: 2px solid #FFD700; padding-bottom: 5px;'>🔬 Especializaciones</h4>", unsafe_allow_html=True)
                                    if df_especi.empty:
                                        st.markdown("<p style='color:#DDDDDD;'>No registradas.</p>", unsafe_allow_html=True)
                                    else:
                                        for _, row in df_especi.iterrows():
                                            inst = get_val(row, ['institucion educativa', 'INSTITUCION EDUCATIVA', 'institución educativa'])
                                            mencion = get_val(row, ['mencion (especialidad / carrera / etc)', 'MENCION (ESPECIALIDAD / CARRERA / ETC)', 'mención'])
                                            anio = get_val(row, ['AÑO', 'año'])
                                            horas = get_val(row, ['horas academicas', 'HORAS ACADEMICAS', 'horas académicas'])
                                            
                                            st.markdown(f"""
                                            <div style='background-color: #FFFFFF; padding: 15px; border-radius: 8px; border: 1px solid #CCCCCC; border-left: 6px solid #9C27B0; margin-bottom: 10px;'>
                                                <div style='margin-bottom: 5px; color: #000000; font-size: 1em; font-weight: bold;'>{mencion}</div>
                                                <div style='margin: 2px 0; color: #000000;'><strong>Institución:</strong> {inst}</div>
                                                <div style='margin: 2px 0; color: #000000;'><strong>Horas Ac.:</strong> {horas} hrs | <strong>Año:</strong> {anio}</div>
                                            </div>
                                            """, unsafe_allow_html=True)

                                with col_der_acad:
                                    # --- 4. DIPLOMADOS ---
                                    st.markdown("<h4 style='color: #FFD700; font-weight: bold; margin-top: 15px; border-bottom: 2px solid #FFD700; padding-bottom: 5px;'>🏅 Diplomados</h4>", unsafe_allow_html=True)
                                    if df_diplo.empty:
                                        st.markdown("<p style='color:#DDDDDD;'>No registrados.</p>", unsafe_allow_html=True)
                                    else:
                                        for _, row in df_diplo.iterrows():
                                            inst = get_val(row, ['institucion educativa', 'INSTITUCION EDUCATIVA', 'institución educativa'])
                                            mencion = get_val(row, ['mencion (especialidad / carrera / etc)', 'MENCION (ESPECIALIDAD / CARRERA / ETC)', 'mención'])
                                            anio = get_val(row, ['AÑO', 'año'])
                                            horas = get_val(row, ['horas academicas', 'HORAS ACADEMICAS', 'horas académicas'])
                                            
                                            st.markdown(f"""
                                            <div style='background-color: #FFFFFF; padding: 15px; border-radius: 8px; border: 1px solid #CCCCCC; border-left: 6px solid #03A9F4; margin-bottom: 10px;'>
                                                <div style='margin-bottom: 5px; color: #000000; font-size: 1em; font-weight: bold;'>{mencion}</div>
                                                <div style='margin: 2px 0; color: #000000;'><strong>Institución:</strong> {inst}</div>
                                                <div style='margin: 2px 0; color: #000000;'><strong>Horas Ac.:</strong> {horas} hrs | <strong>Año:</strong> {anio}</div>
                                            </div>
                                            """, unsafe_allow_html=True)

                                    # --- 5. CURSOS ---
                                    st.markdown("<h4 style='color: #FFD700; font-weight: bold; margin-top: 15px; border-bottom: 2px solid #FFD700; padding-bottom: 5px;'>📚 Cursos</h4>", unsafe_allow_html=True)
                                    if df_cursos.empty:
                                        st.markdown("<p style='color:#DDDDDD;'>No registrados.</p>", unsafe_allow_html=True)
                                    else:
                                        for _, row in df_cursos.iterrows():
                                            inst = get_val(row, ['institucion educativa', 'INSTITUCION EDUCATIVA', 'institución educativa'])
                                            mencion = get_val(row, ['mencion (especialidad / carrera / etc)', 'MENCION (ESPECIALIDAD / CARRERA / ETC)', 'mención'])
                                            anio = get_val(row, ['AÑO', 'año'])
                                            horas = get_val(row, ['horas academicas', 'HORAS ACADEMICAS', 'horas académicas'])
                                            
                                            st.markdown(f"""
                                            <div style='background-color: #FFFFFF; padding: 15px; border-radius: 8px; border: 1px solid #CCCCCC; border-left: 6px solid #4CAF50; margin-bottom: 10px;'>
                                                <div style='margin-bottom: 5px; color: #000000; font-size: 1em; font-weight: bold;'>{mencion}</div>
                                                <div style='margin: 2px 0; color: #000000;'><strong>Institución:</strong> {inst}</div>
                                                <div style='margin: 2px 0; color: #000000;'><strong>Horas Ac.:</strong> {horas} hrs | <strong>Año:</strong> {anio}</div>
                                            </div>
                                            """, unsafe_allow_html=True)
                                # --- TABLA DE SELECCIÓN PARA EDITAR/ELIMINAR ---
                                st.markdown("<br>", unsafe_allow_html=True)
                                with st.expander("⚙️ Clic aquí para Editar o Eliminar Formación Académica"):
                                    st.markdown("<span style='color:#A0A0A0; font-size:14px;'>Activa la casilla <b>SEL</b> en la tabla de abajo para modificar o eliminar un registro.</span>", unsafe_allow_html=True)
                                    
                                    # Traemos los datos del trabajador para esta pestaña
                                    df_fa = dfs[h_name][dfs[h_name]["dni"].astype(str) == str(dni_buscado)].copy()
                                    
                                    if not df_fa.empty:
                                        # Agregamos la columna SEL al inicio
                                        df_fa.insert(0, "SEL", False)
                                        # Mostramos la tabla interactiva
                                        ed = st.data_editor(
                                            df_fa,
                                            hide_index=True,
                                            use_container_width=True,
                                            disabled=[c for c in df_fa.columns if c != "SEL"],
                                            key="editor_form_acad" # Llave única importante
                                        )
                                        # Capturamos la fila que el usuario seleccione para que el bloque de abajo la lea
                                        sel = ed[ed["SEL"] == True]
                                    else:
                                        st.info("No hay registros para mostrar.")
                                        sel = pd.DataFrame()
                            
                            # ==========================================
                            # PESTAÑA: DATOS FAMILIARES
                            # ==========================================
                            elif h_name == "DATOS FAMILIARES":
                                st.markdown("<h3 style='color: #FFD700; margin-bottom: 20px;'>👨‍👩‍👧‍👦 Datos Familiares</h3>", unsafe_allow_html=True)
                                
                                # --- 1. BUSCAR LA DIRECCIÓN DEL TRABAJADOR ---
                                dir_trabajador = ""
                                if not dfs["DATOS GENERALES"].empty:
                                    df_gen = dfs["DATOS GENERALES"]
                                    # Buscamos la fila del trabajador por su DNI
                                    datos_trabajador = df_gen[df_gen["dni"].astype(str) == str(dni_buscado)]
                                    if not datos_trabajador.empty:
                                        # Buscamos exactamente la columna sin importar si panda la lee en mayúscula o minúscula
                                        for col in datos_trabajador.columns:
                                            if str(col).strip().upper() == "DIRECCION":
                                                val = datos_trabajador.iloc[0][col]
                                                if pd.notna(val) and str(val).strip() != "":
                                                    dir_trabajador = str(val)
                                                break

                                # --- 2. MOSTRAR FAMILIARES REGISTRADOS (CON TODOS LOS DATOS) ---
                                st.markdown("<h4 style='color: #FFD700; border-bottom: 2px solid #FFD700; padding-bottom: 5px;'>📋 Familiares Registrados</h4>", unsafe_allow_html=True)
                                
                                def get_fam_val(r, col_name):
                                    for col in r.index:
                                        if str(col).strip().lower() == col_name.lower():
                                            val = r[col]
                                            if pd.notna(val) and str(val).strip() != "":
                                                return str(val)
                                    return "-"

                                if len(vst) == 0:
                                    st.markdown("<p style='color:#DDDDDD;'>No hay familiares registrados aún.</p>", unsafe_allow_html=True)
                                else:
                                    for idx, row in vst.iterrows():
                                        f_dni = get_fam_val(row, "dni familiar")
                                        if f_dni == "-": f_dni = get_fam_val(row, "dni_familiar") # Por si acaso
                                        f_parentesco = get_fam_val(row, "parentesco")
                                        f_nombres = get_fam_val(row, "nombres y apellidos")
                                        f_edad = get_fam_val(row, "edad")
                                        f_estado = get_fam_val(row, "estado")
                                        f_celular = get_fam_val(row, "celular")
                                        f_correo = get_fam_val(row, "correo")
                                        f_domicilio = get_fam_val(row, "domicilio")
                                        f_sit_acad = get_fam_val(row, "situacion academica")
                                        f_emergencia = get_fam_val(row, "contacto emergencia").lower()
                                        
                                        badge_emergencia = "<span style='color: #FF5252; font-size: 0.9em;'>🚨 <b>CONTACTO DE EMERGENCIA</b></span>" if f_emergencia in ["sí", "si", "true", "1"] else ""
                                        
                                        st.markdown(f"""
                                        <div style='background-color: #FFFFFF; padding: 15px; border-radius: 8px; border: 1px solid #CCCCCC; border-left: 6px solid #2196F3; margin-bottom: 10px; color: #000000;'>
                                            <div style='margin-bottom: 10px; font-size: 1.1em; font-weight: bold; border-bottom: 1px solid #EEEEEE; padding-bottom: 5px;'>
                                                {f_nombres} <span style='color: #666666; font-size: 0.9em;'>({f_parentesco})</span> {badge_emergencia}
                                            </div>
                                            <div style='display: grid; grid-template-columns: 1fr 1fr; gap: 8px; font-size: 0.95em;'>
                                                <div><strong>DNI:</strong> {f_dni}</div>
                                                <div><strong>Edad:</strong> {f_edad} años</div>
                                                <div><strong>Estado:</strong> {f_estado}</div>
                                                <div><strong>Celular:</strong> {f_celular}</div>
                                                <div><strong>Correo:</strong> {f_correo}</div>
                                                <div><strong>Sit. Académica:</strong> {f_sit_acad}</div>
                                                <div style='grid-column: span 2;'><strong>Domicilio:</strong> {f_domicilio}</div>
                                            </div>
                                        </div>
                                        """, unsafe_allow_html=True)

                                st.markdown("<br>", unsafe_allow_html=True)

                                # --- 3. TABLA DESPLEGABLE PARA EDICIÓN ---
                                with st.expander("⚙️ Clic aquí para Editar o Eliminar un Familiar"):
                                    st.markdown("<p style='color:#DDDDDD;'>Activa la casilla <b>SEL</b> en la tabla de abajo para modificar o eliminar un registro.</p>", unsafe_allow_html=True)
                                    
                                    # ✅ AYUDA VISUAL PARA COPIAR Y PEGAR LA DIRECCIÓN
                                    if dir_trabajador:
                                        st.info(f"💡 **Tip para la edición:** Si el familiar vive con el trabajador, simplemente copia y pega esta dirección en la tabla:  \n**{dir_trabajador}**")
                                        
                                    st.markdown("""<style>[data-testid="stDataEditor"] { border: 2px solid #FFD700 !important; border-radius: 8px !important; }</style>""", unsafe_allow_html=True)
                                    ed = st.data_editor(vst, hide_index=True, use_container_width=True, column_config=col_conf, key=f"ed_{h_name}_oculta")
                                    sel = ed[ed["SEL"] == True]

                                # --- 4. FORMULARIO DENTRO DE "NUEVO REGISTRO" ---
                                st.markdown("<br>", unsafe_allow_html=True)
                                with st.expander("➕ NUEVO REGISTRO"):
                                    
                                    st.markdown("<p style='color:#DDDDDD; font-style: italic;'>Rellena los datos para agregar un familiar.</p>", unsafe_allow_html=True)
                                    col_f1, col_f2 = st.columns(2)
                                    
                                    with col_f1:
                                        parentesco = st.selectbox("Parentesco", ["Cónyuge / Conviviente", "Hijo(a)", "Madre", "Padre", "Hermano(a)", "Familiar Adicional (Otros)"])
                                        dni_fam = st.text_input("DNI del Familiar", max_chars=8)
                                        
                                        if dni_fam and len(dni_fam) >= 8:
                                            if not dfs["DATOS GENERALES"].empty:
                                                es_trabajador = dfs["DATOS GENERALES"][dfs["DATOS GENERALES"]["dni"].astype(str) == str(dni_fam)]
                                                if not es_trabajador.empty:
                                                    nombre_vinculo = es_trabajador.iloc[0].get("apellidos y nombres", "Trabajador")
                                                    st.success(f"🔗 ¡Vínculo detectado! Este familiar es trabajador activo: **{nombre_vinculo}**")
                                        
                                        nombres_fam = st.text_input("Apellidos y Nombres")
                                        
                                        f_nac_fam = st.date_input("Fecha de Nacimiento", min_value=date(1920, 1, 1), max_value=date.today())
                                        hoy = date.today()
                                        edad_fam = hoy.year - f_nac_fam.year - ((hoy.month, hoy.day) < (f_nac_fam.month, f_nac_fam.day))
                                        st.info(f"🎂 Edad calculada: **{edad_fam} años**")

                                    with col_f2:
                                        estado_fam = st.selectbox("Estado", ["Vivo", "Fallecido", "Otra condición"])
                                        
                                        if estado_fam == "Vivo":
                                            cel_fam = st.text_input("Celular")
                                            correo_fam = st.text_input("Correo Electrónico")
                                            
                                            # --- CHECK DE VIVE CON EL TRABAJADOR ---
                                            st.markdown("---")
                                            vive_juntos = st.checkbox("🏠 Vive con el trabajador")
                                            
                                            if vive_juntos:
                                                domicilio_fam = st.text_input("Domicilio del familiar", value=dir_trabajador, key="domicilio_juntos")
                                                if not dir_trabajador:
                                                    st.warning("⚠️ Ojo: El trabajador no tiene una dirección registrada en su pestaña de Datos Generales.")
                                            else:
                                                domicilio_fam = st.text_input("Domicilio del familiar", value="", key="domicilio_separado")
                                        else:
                                            cel_fam = "-"
                                            correo_fam = "-"
                                            domicilio_fam = "-"
                                            vive_juntos = False

                                        sit_acad_fam = st.selectbox("Situación Académica", [
                                            "Ninguna / No aplica",
                                            "Estudiando Primaria", "Estudiando Secundaria", "Estudiando Superior",
                                            "Estudios Concluidos Primaria", "Estudios Concluidos Secundaria", "Estudios Concluidos Superior"
                                        ])
                                        
                                        contacto_emergencia = st.checkbox("🚨 Es Contacto de Emergencia Principal")

                                    if st.button("💾 Guardar Familiar", type="primary"):
                                        if not dni_fam or not nombres_fam:
                                            st.error("⚠️ El DNI y los Nombres son obligatorios.")
                                        else:
                                            new_row = {
                                                "dni": str(dni_buscado),
                                                "dni familiar": str(dni_fam),
                                                "parentesco": parentesco,
                                                "nombres y apellidos": nombres_fam,
                                                "fecha de nacimiento": str(f_nac_fam),
                                                "edad": edad_fam,
                                                "domicilio": domicilio_fam,
                                                "estado": estado_fam,
                                                "celular": cel_fam,
                                                "correo": correo_fam,
                                                "situacion academica": sit_acad_fam,
                                                "contacto emergencia": "Sí" if contacto_emergencia else "No"
                                            }
                                            
                                            if not dfs[h_name].empty and "id" in dfs[h_name].columns:
                                                new_row["id"] = dfs[h_name]["id"].max() + 1
                                            elif "id" in dfs[h_name].columns:
                                                new_row["id"] = 1
                                                
                                            dfs[h_name] = pd.concat([dfs[h_name], pd.DataFrame([new_row])], ignore_index=True)
                                            save_data(dfs)
                                            st.success("✅ Familiar guardado correctamente.")
                                            st.rerun()                   
                        # ==========================================
                        # BOTÓN DE IMPRESIÓN DE PAPELETA (SOLO EN VACACIONES)
                        # ==========================================
                        if h_name == "VACACIONES" and not sel.empty:
                            st.markdown("---")
                            # Necesitamos el cargo actual y fecha de ingreso del trabajador
                            current_cargo = "TRABAJADOR" # Default
                            f_ingreso_val = ""
                            df_c_data = dfs["CONTRATOS"][dfs["CONTRATOS"]["dni"] == dni_buscado]
                            
                            if not df_c_data.empty:
                                try:
                                    # Obtener cargo (último contrato)
                                    last_contract = df_c_data.assign(f_fin_dt=pd.to_datetime(df_c_data['f_fin'], errors='coerce')).sort_values('f_fin_dt').iloc[-1]
                                    current_cargo = last_contract.get("cargo", "TRABAJADOR")
                                    
                                    # Obtener fecha de ingreso (primer contrato de planilla)
                                    df_planilla = df_c_data[df_c_data["tipo contrato"].astype(str).str.lower().str.contains("planilla", na=False)]
                                    if not df_planilla.empty:
                                        f_min = pd.to_datetime(df_planilla['f_inicio'], errors='coerce').min()
                                        if pd.notnull(f_min): f_ingreso_val = f_min.date()
                                except: pass

                            # Capturar datos de la fila seleccionada
                            r_sel = sel.iloc[0]
                            
                            # 1. BÚSQUEDA INTELIGENTE DE COLUMNAS (Barre todas las coincidencias)
                            cols_per = [c for c in r_sel.index if "PERIODO" in str(c).upper()]
                            cols_ini = [c for c in r_sel.index if "INICIO" in str(c).upper()]
                            cols_fin = [c for c in r_sel.index if "FIN" in str(c).upper()]
                            cols_dias = [c for c in r_sel.index if "GOZADOS" in str(c).upper()]

                            # Función para ignorar los 'NaT' y sacar el valor real
                            def get_valid_val(cols):
                                for c in cols:
                                    val = r_sel.get(c)
                                    if pd.notnull(val) and str(val).strip() not in ["", "NaT", "None"]:
                                        return val
                                return None

                            p_papeleta = str(get_valid_val(cols_per) or "")
                            fi_papeleta_raw = get_valid_val(cols_ini)
                            ff_papeleta_raw = get_valid_val(cols_fin)
                            dg_papeleta_raw = get_valid_val(cols_dias) or 0
                            
                            # 2. CONVERSIÓN ULTRA ROBUSTA DE FECHAS
                            try:
                                fi_papeleta = pd.to_datetime(fi_papeleta_raw).date() if fi_papeleta_raw else None
                            except:
                                fi_papeleta = None
                                
                            try:
                                ff_papeleta = pd.to_datetime(ff_papeleta_raw).date() if ff_papeleta_raw else None
                            except:
                                ff_papeleta = None

                            # 3. LIMPIEZA DE LOS DÍAS GOZADOS
                            try:
                                dg_papeleta = int(float(dg_papeleta_raw))
                            except:
                                dg_papeleta = 0

                            # 4. BOTÓN Y VALIDACIÓN FINAL
                            if st.button(f"📄 Generar Papeleta de Impresión (Periodo {p_papeleta})", key="btn_print_vaca_tab", use_container_width=False):
                                if fi_papeleta is None or ff_papeleta is None:
                                    st.error(f"⚠️ Aún no se detectan fechas válidas. Inicio extraído: '{fi_papeleta_raw}' | Fin extraído: '{ff_papeleta_raw}'")
                                else:
                                    # AQUÍ LLAMAMOS A LA FUNCIÓN CON TODOS LOS DATOS
                                    papeleta_word = gen_papeleta_vac(ape_c, nom_p_c, dni_buscado, current_cargo, f_ingreso_val, p_papeleta, fi_papeleta, ff_papeleta, dg_papeleta)
                                    if papeleta_word:
                                        st.markdown("""<style>[data-testid="stDownloadButton"] button { background-color: #FFD700 !important; color: #4A0000 !important; font-weight: bold !important; border: 2px solid #4A0000 !important; width: 100% !important; }</style>""", unsafe_allow_html=True)
                                        st.download_button(
                                            label=f"⬇️ Descargar Papeleta - {nom_c}.docx",
                                            data=papeleta_word,
                                            file_name=f"Papeleta_{dni_buscado}_{p_papeleta}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            key="dl_papeleta_tab"
                                        )
                            st.markdown("---")
                        
                        if not es_lector:
                            # 1. Validamos que la hoja se haya cargado correctamente en 'dfs'
                            if h_name not in dfs:
                                st.error(f"⚠️ Error crítico: No se pudo cargar la pestaña '{h_name}'. Por favor, entra al Google Sheets y elimina las columnas duplicadas.")
                            else:
                                # 2. Si la hoja existe, ejecutamos el código normalmente
                                col_a, col_b = st.columns(2)
                                cols_reales = [c for c in dfs[h_name].columns if c.lower() not in ["id", "dni", "apellidos y nombres", "apellidos", "nombres"]]
                                
                                with col_a:
                                    df_filtro = dfs[h_name][dfs[h_name]["dni"] == dni_buscado] if not dfs[h_name].empty else pd.DataFrame()
                                    
                                    if h_name == "DATOS GENERALES" and len(df_filtro) > 0:
                                        st.info("📌 Los datos generales ya están registrados. Selecciona el registro en la tabla de arriba para editarlos.")
                                    elif h_name == "DATOS FAMILIARES":
                                        pass # No dibujamos nada extra porque estas pestañas ya tienen su propio diseño arriba
                                    else:
                                        # ¡AQUÍ ESTÁ EL CAMBIO IMPORTANTE!
                                        with st.expander("➕ Nuevo Registro"):
                                            
                                            # ==========================================
                                            # NUEVO REGISTRO REACTIVO DE VACACIONES
                                            # ==========================================
                                            if h_name == "VACACIONES":
                                                st.markdown("<div style='font-size: 1.5em; font-weight: bold; color: white; background-color: #4A0000; padding: 10px; border-radius: 8px; margin-bottom: 15px;'>➕ Registrar Nuevas Vacaciones</div>", unsafe_allow_html=True)
                                                
                                                if detalles:
                                                    opciones_periodo = [d["Periodo"] for d in detalles]
                                                    dict_generados = {d["Periodo"]: d["Días Generados"] for d in detalles}
                                                    dict_saldo_actual = {d["Periodo"]: d["Saldo"] for d in detalles}
                                                else:
                                                    opciones_periodo = ["Sin periodo calculado"]
                                                    dict_generados = {"Sin periodo calculado": 0}
                                                    dict_saldo_actual = {"Sin periodo calculado": 0}

                                                sel_periodo = st.selectbox("Periodo Vacacional", options=opciones_periodo)
                                                
                                                col_f1, col_f2 = st.columns(2)
                                                with col_f1:
                                                    f_ini_val = st.date_input("Fecha de Salida (Inicio)", min_value=date(1950, 1, 1), max_value=date(2100, 12, 31))
                                                with col_f2:
                                                    f_fin_val = st.date_input("Fecha de Retorno (Último día)", min_value=date(1950, 1, 1), max_value=date(2100, 12, 31))

                                                dias_gozar_calc = 0
                                                if f_fin_val >= f_ini_val:
                                                    dias_gozar_calc = (f_fin_val - f_ini_val).days + 1
                                                
                                                gen_periodo = dict_generados.get(sel_periodo, 0)
                                                saldo_previo = dict_saldo_actual.get(sel_periodo, 0)
                                                nuevo_saldo = saldo_previo - dias_gozar_calc

                                                # Lógica de colores para el saldo
                                                if nuevo_saldo < 0:
                                                    txt_saldo = f":red[{nuevo_saldo:.2f} (¡Saldo Negativo!)]"
                                                elif nuevo_saldo == 0:
                                                    txt_saldo = f"{nuevo_saldo:.2f}"
                                                else:
                                                    txt_saldo = f":green[{nuevo_saldo:.2f}]"
                                                
                                                st.markdown(f"""
                                                **Resumen:**
                                                * **Días a Gozar (Calculado):** {dias_gozar_calc}
                                                * **Saldo Restante:** {txt_saldo}
                                                """)
                                                
                                                if st.button("💾 Guardar Registro de Vacaciones", type="primary", use_container_width=False):
                                                    if dias_gozar_calc <= 0:
                                                        st.error("⚠️ La Fecha de Fin debe ser igual o posterior a la Fecha de Inicio.")
                                                    else:
                                                        new_row = {"DNI": dni_buscado, "PERIODO": sel_periodo, "F_INICIO": f_ini_val, "F_FIN": f_fin_val, "DIAS GOZADOS": dias_gozar_calc}
                                                        if not dfs[h_name].empty and "id" in dfs[h_name].columns:
                                                            new_row["id"] = dfs[h_name]["id"].max() + 1
                                                        elif "id" in dfs[h_name].columns:
                                                            new_row["id"] = 1
                                                        
                                                        dfs[h_name] = pd.concat([dfs[h_name], pd.DataFrame([new_row])], ignore_index=True)
                                                        save_data(dfs)
                                                        st.session_state['just_saved_vacation'] = new_row
                                                        st.success("✅ Registro guardado correctamente.")
                                                        st.rerun()

                                            # ==========================================
                                            # FORMULARIOS NORMALES PARA EL RESTO DE HOJAS
                                            # ==========================================
                                            # ==========================================
                                            # NUEVO REGISTRO: EXPERIENCIA LABORAL
                                            # ==========================================
                                            elif h_name == "EXP. LABORAL":
                                                st.markdown("<div style='font-size: 1.5em; font-weight: bold; color: white; background-color: #004A80; padding: 10px; border-radius: 8px; margin-bottom: 15px;'>💼 Registrar Experiencia Externa</div>", unsafe_allow_html=True)
                                                
                                                col_e1, col_e2 = st.columns(2)
                                                with col_e1:
                                                    tipo_exp = st.selectbox("Tipo de Experiencia", ["Administrativo", "Docente"])
                                                    lugar_exp = st.text_input("Lugar (Empresa/Institución)")
                                                    puesto_exp = st.text_input("Puesto o Cargo ocupado")
                                                with col_e2:
                                                    # --- AQUÍ APLICAMOS LOS LÍMITES DE AÑOS ---
                                                    f_ini_exp = st.date_input("Fecha de Inicio", min_value=date(1950, 1, 1), max_value=date(2100, 12, 31))
                                                    f_fin_exp = st.date_input("Fecha de Fin", min_value=date(1950, 1, 1), max_value=date(2100, 12, 31))
                                                    # ------------------------------------------
                                                    motivo_exp = st.selectbox("Motivo de Cese", MOTIVOS_CESE)
                                                    
                                                if st.button("💾 Guardar Experiencia", type="primary", use_container_width=False):
                                                    if f_fin_exp < f_ini_exp:
                                                        st.error("⚠️ La Fecha de Fin no puede ser anterior a la Fecha de Inicio.")
                                                    elif not lugar_exp or not puesto_exp:
                                                        st.error("⚠️ El Lugar y el Puesto son campos obligatorios.")
                                                    else:
                                                        new_row = {
                                                            "dni": str(dni_buscado), 
                                                            "tipo de experiencia": tipo_exp, 
                                                            "lugar": lugar_exp, 
                                                            "puesto": puesto_exp, 
                                                            "fecha de inicio": f_ini_exp, 
                                                            "fecha de fin": f_fin_exp, 
                                                            "motivo de cese": motivo_exp
                                                        }
                                                        if not dfs[h_name].empty and "id" in dfs[h_name].columns:
                                                            new_row["id"] = dfs[h_name]["id"].max() + 1
                                                        elif "id" in dfs[h_name].columns:
                                                            new_row["id"] = 1
                                                            
                                                        dfs[h_name] = pd.concat([dfs[h_name], pd.DataFrame([new_row])], ignore_index=True)
                                                        save_data(dfs)
                                                        st.success("✅ Experiencia guardada correctamente.")
                                                        st.rerun()
                                           # --- CARGA MAESTRA DE PARÁMETROS PARA ESTAS PESTAÑAS ---
                                            df_para = dfs.get("PARAMETROS", pd.DataFrame())
                                            if not df_para.empty:
                                                df_para.columns = df_para.columns.str.strip().str.replace(" ", "_").str.upper()
                                            
                                            def get_lista(columna, default):
                                                if columna in df_para.columns:
                                                    lista = df_para[columna].dropna().astype(str).str.strip().unique().tolist()
                                                    lista = [i for i in lista if i and i.lower() != "nan"]
                                                    return lista if lista else default
                                                return default
                                            # --------------------------------------------------------

                                           # ==========================================
                                            # NUEVO REGISTRO: FORMACIÓN ACADÉMICA (DINÁMICO)
                                            # ==========================================
                                            elif h_name == "FORM. ACADEMICA":
                                                # --- CARGA MAESTRA DE PARÁMETROS ---
                                                df_para = dfs.get("PARAMETROS", pd.DataFrame())
                                                if not df_para.empty:
                                                    df_para.columns = df_para.columns.str.strip().str.replace(" ", "_").str.upper()
                                                
                                                def get_lista(columna, default):
                                                    if columna in df_para.columns:
                                                        lista = df_para[columna].dropna().astype(str).str.strip().unique().tolist()
                                                        lista = [i for i in lista if i and i.lower() != "nan"]
                                                        return lista if lista else default
                                                    return default
                                                # -----------------------------------

                                                # Listas dinámicas para Formación
                                                lst_tipo_estudio = get_lista("TIPO_ESTUDIO", ["Grados y Títulos", "Estudios Terminados", "Especializaciones", "Diplomados", "Cursos"])
                                                lst_grado = get_lista("GRADO_OBTENIDO", ["Bachiller", "Título Profesional", "Magíster", "Doctor"])
                                                lst_estado_est = get_lista("ESTADO_ESTUDIO", ["Concluido", "Estudiando", "En abandono"])

                                                st.markdown("""
                                                <div style='color: #000000; font-size: 1.5em; font-weight: bold; margin-bottom: 5px;'>🎓 Registrar Nuevo Estudio</div>
                                                <div style='color: #000000; margin-bottom: 15px;'>Selecciona el tipo de estudio para ver los campos requeridos.</div>
                                                """, unsafe_allow_html=True)
                                                
                                                tipo_estudio = st.selectbox("📌 Tipo de Estudio", lst_tipo_estudio)

                                                grado = "N/A"
                                                institucion = ""
                                                mencion = ""
                                                anio = ""
                                                estado = "N/A"
                                                horas = "N/A"

                                                col_f1, col_f2 = st.columns(2)

                                                if tipo_estudio in ["Grados y Títulos", "Pregrado", "Maestría", "Doctorado"]:
                                                    with col_f1:
                                                        grado = st.selectbox("Grado o Título Obtenido", lst_grado)
                                                        institucion = st.text_input("Institución Educativa")
                                                    with col_f2:
                                                        mencion = st.text_input("Mención (Especialidad / Carrera)")
                                                        anio = st.text_input("Año")

                                                elif tipo_estudio == "Estudios Terminados (Sin grado) o Inconclusos":
                                                    with col_f1:
                                                        institucion = st.text_input("Institución Educativa")
                                                        mencion = st.text_input("Mención (Especialidad / Carrera)")
                                                    with col_f2:
                                                        anio = st.text_input("Año")
                                                        estado = st.selectbox("Estado", lst_estado_est)

                                                else: # Especializaciones, Diplomados, Cursos, etc.
                                                    with col_f1:
                                                        institucion = st.text_input("Institución Educativa")
                                                        mencion = st.text_input("Mención (Nombre del estudio)")
                                                    with col_f2:
                                                        horas = st.text_input("Horas Académicas")
                                                        anio = st.text_input("Año")

                                                if st.button("💾 Guardar Estudio", type="primary", use_container_width=False):
                                                    if not institucion or not mencion:
                                                        st.error("⚠️ La Institución y la Mención son campos obligatorios.")
                                                    else:
                                                        new_row = {
                                                            "dni": str(dni_buscado), 
                                                            "tipo de estudio": tipo_estudio, 
                                                            "grado o titulo obtenido": grado,
                                                            "institucion educativa": institucion, 
                                                            "mencion (especialidad / carrera / etc)": mencion, 
                                                            "año": anio, 
                                                            "estado": estado, 
                                                            "horas academicas": horas
                                                        }
                                                        
                                                        if not dfs[h_name].empty and "id" in dfs[h_name].columns:
                                                            new_row["id"] = dfs[h_name]["id"].max() + 1
                                                        elif "id" in dfs[h_name].columns:
                                                            new_row["id"] = 1
                                                            
                                                        dfs[h_name] = pd.concat([dfs[h_name], pd.DataFrame([new_row])], ignore_index=True)
                                                        save_data(dfs)
                                                        st.success("✅ Estudio guardado correctamente.")
                                                        st.rerun()

                                            # ==========================================
                                            # NUEVO REGISTRO: INVESTIGACIÓN
                                            # ==========================================
                                            elif h_name == "INVESTIGACION":
                                                # --- CARGA MAESTRA DE PARÁMETROS ---
                                                df_para = dfs.get("PARAMETROS", pd.DataFrame())
                                                if not df_para.empty:
                                                    df_para.columns = df_para.columns.str.strip().str.replace(" ", "_").str.upper()
                                                
                                                def get_lista(columna, default):
                                                    if columna in df_para.columns:
                                                        lista = df_para[columna].dropna().astype(str).str.strip().unique().tolist()
                                                        lista = [i for i in lista if i and i.lower() != "nan"]
                                                        return lista if lista else default
                                                    return default
                                                # -----------------------------------
                                                
                                                # Listas dinámicas para Investigación
                                                lst_nivel_ren = get_lista("NIVEL_RENACYT", ["No tiene", "Nivel VII", "Nivel VI"])
                                                lst_bd = get_lista("BASE_DATOS_INV", ["Scopus", "Web of Science (WoS)", "SciELO"])
                                                lst_cuartil = get_lista("CUARTIL_INV", ["Q1", "Q2", "Q3", "Q4", "No aplica"])
                                                lst_rol_p = get_lista("ROL_PROYECTO", ["Investigador Principal", "Co-Investigador"])
                                                lst_est_p = get_lista("ESTADO_PROYECTO", ["En postulación", "En ejecución", "Finalizado"])
                                                lst_rol_s = get_lista("ROL_SEMILLERO", ["Coordinador/Asesor", "Miembro Investigador"])
                                                lst_est_s = get_lista("ESTADO_SEMILLERO", ["Activo", "Inactivo"])

                                                st.markdown("<div style='font-size: 1.5em; font-weight: bold; color: white; background-color: #4A0000; padding: 10px; border-radius: 8px; margin-bottom: 15px;'>🔬 Registrar Actividad de Investigación</div>", unsafe_allow_html=True)

                                                tipo_registro = st.selectbox(
                                                    "¿Qué tipo de registro deseas ingresar?",
                                                    ["Datos Generales (CTI Vitae / RENACYT)", "Publicación Científica", "Fondo Concursable", "Semillero de Investigación"],
                                                    key="sel_tipo_inv"
                                                )

                                                with st.form("form_nuevo_investigacion", clear_on_submit=True):
                                                    enlace_cti = codigo_renacyt = nivel_renacyt = ""
                                                    titulo_pub = base_datos = nombre_revista = cuartil = doi_url = ""
                                                    anio_pub = date.today().year
                                                    nombre_proy = entidad_fin = rol_proy = estado_proy = ""
                                                    monto_proy = 0.0
                                                    nombre_semillero = resolucion = rol_semillero = estado_semillero = ""
                                                    
                                                    if tipo_registro == "Datos Generales (CTI Vitae / RENACYT)":
                                                        col1, col2, col3 = st.columns(3)
                                                        with col1: enlace_cti = st.text_input("Enlace CTI Vitae")
                                                        with col2: codigo_renacyt = st.text_input("Código RENACYT")
                                                        with col3: nivel_renacyt = st.selectbox("Nivel RENACYT", lst_nivel_ren)
                                                        
                                                    elif tipo_registro == "Publicación Científica":
                                                        titulo_pub = st.text_input("Título de la Publicación")
                                                        col1, col2 = st.columns(2)
                                                        with col1:
                                                            base_datos = st.selectbox("Base de Datos", lst_bd)
                                                            cuartil = st.selectbox("Cuartil", lst_cuartil)
                                                            anio_pub = st.number_input("Año de Publicación", min_value=1950, max_value=2100, step=1, value=date.today().year)
                                                        with col2:
                                                            nombre_revista = st.text_input("Nombre de la Revista")
                                                            doi_url = st.text_input("DOI / URL del artículo")
                                                            
                                                    elif tipo_registro == "Fondo Concursable":
                                                        nombre_proy = st.text_input("Nombre del Proyecto")
                                                        col1, col2 = st.columns(2)
                                                        with col1:
                                                            entidad_fin = st.text_input("Entidad Financiadora")
                                                            monto_proy = st.number_input("Monto Adjudicado (S/.)", min_value=0.0, step=100.0)
                                                        with col2:
                                                            rol_proy = st.selectbox("Rol en el Proyecto", lst_rol_p)
                                                            estado_proy = st.selectbox("Estado del Proyecto", lst_est_p)
                                                            
                                                    elif tipo_registro == "Semillero de Investigación":
                                                        nombre_semillero = st.text_input("Nombre del Semillero")
                                                        col1, col2 = st.columns(2)
                                                        with col1:
                                                            resolucion = st.text_input("Resolución de Creación")
                                                            rol_semillero = st.selectbox("Rol en el Semillero", lst_rol_s)
                                                        with col2:
                                                            estado_semillero = st.selectbox("Estado", lst_est_s)
                                                            
                                                    if st.form_submit_button("💾 Guardar Registro de Investigación"):
                                                        nid = dfs[h_name]["id"].max() + 1 if not dfs[h_name].empty and "id" in dfs[h_name].columns else 1
                                                        
                                                        nuevo_registro = {
                                                            "id": nid, "dni": str(dni_buscado), "tipo de registro": tipo_registro,
                                                            "enlace cti vitae": enlace_cti, "codigo renacyt": codigo_renacyt, "nivel renacyt": nivel_renacyt,
                                                            "titulo de publicacion": titulo_pub, "base de datos": base_datos, "nombre de revista": nombre_revista,
                                                            "cuartil": cuartil, "año de publicacion": anio_pub, "doi o url": doi_url,
                                                            "nombre del proyecto": nombre_proy, "entidad financiadora": entidad_fin,
                                                            "rol en el proyecto": rol_proy, "monto adjudicado": monto_proy, "estado del proyecto": estado_proy,
                                                            "nombre del semillero": nombre_semillero, "resolucion": resolucion, "rol en el semillero": rol_semillero,
                                                            "estado del semillero": estado_semillero
                                                        }
                                                        
                                                        dfs[h_name] = pd.concat([dfs[h_name], pd.DataFrame([nuevo_registro])], ignore_index=True)
                                                        save_data(dfs)
                                                        st.success("✅ ¡Registro de investigación guardado exitosamente!")
                                                        st.rerun()

                                            # ==========================================
                                            # CONTRATOS Y CONDICIONES LABORALES
                                            # ==========================================
                                            else:
                                                # --- CARGA MAESTRA DE PARÁMETROS ---
                                                df_para = dfs.get("PARAMETROS", pd.DataFrame())
                                                if not df_para.empty:
                                                    df_para.columns = df_para.columns.str.strip().str.replace(" ", "_").str.upper()
                                                
                                                def get_lista(columna, default):
                                                    if columna in df_para.columns:
                                                        lista = df_para[columna].dropna().astype(str).str.strip().unique().tolist()
                                                        lista = [i for i in lista if i and i.lower() != "nan"]
                                                        return lista if lista else default
                                                    return default
                                                # -----------------------------------

                                                # Listas dinámicas para Contratos
                                                lst_cargo = get_lista("CARGO", ["Docente", "Administrativo"])
                                                lst_area = get_lista("AREA", ["Recursos Humanos", "Dirección Académica"])
                                                
                                                es_renovacion = False
                                                if h_name == "CONTRATOS" and not df_contratos.empty:
                                                    es_renovacion = st.checkbox("🔄 Es Renovación (Copiar datos del último contrato)")
                                                
                                                with st.form(f"f_add_{h_name}", clear_on_submit=True):
                                                    if h_name == "CONTRATOS":
                                                        d_car = ""
                                                        d_area = ""  
                                                        d_rem = 0.0
                                                        d_bon = ""
                                                        d_cond = ""
                                                        d_ini = date.today()
                                                        d_fin = date.today()
                                                        d_ttrab = "Administrativo"
                                                        d_mod = "Presencial"
                                                        d_temp = "Plazo fijo"
                                                        d_tcont = "Planilla completo"
                                                        
                                                        if es_renovacion and not df_contratos.empty:
                                                            last_c = df_contratos.assign(f_fin_dt=pd.to_datetime(df_contratos['f_fin'], errors='coerce')).sort_values('f_fin_dt').iloc[-1]
                                                            d_car = str(last_c.get("cargo", ""))
                                                            d_area = str(last_c.get("area", ""))
                                                            try: d_rem = float(last_c.get("remuneracion basica", 0.0))
                                                            except: pass
                                                            d_bon = str(last_c.get("bonificacion", ""))
                                                            d_cond = str(last_c.get("condicion de trabajo", ""))
                                                            try: d_ini = pd.to_datetime(last_c["f_fin"]).date() + pd.Timedelta(days=1)
                                                            except: pass
                                                            
                                                            v_tt = str(last_c.get("tipo de trabajador", ""))
                                                            if v_tt in ["Administrativo", "Docente", "Externo"]: d_ttrab = v_tt
                                                            v_m = str(last_c.get("modalidad", ""))
                                                            if v_m in ["Presencial", "Semipresencial", "Virtual"]: d_mod = v_m
                                                            v_te = str(last_c.get("temporalidad", ""))
                                                            if v_te in ["Plazo fijo", "Plazo indeterminado", "Ordinarizado"]: d_temp = v_te
                                                            v_tc = str(last_c.get("tipo contrato", ""))
                                                            if v_tc in ["Planilla completo", "Tiempo Parcial", "Recibo por Honorarios", "Otro"]: d_tcont = v_tc

                                                        # Lógica inteligente para pre-seleccionar Cargo y Área al renovar
                                                        if d_car and d_car not in lst_cargo: lst_cargo.append(d_car)
                                                        idx_car = lst_cargo.index(d_car) if d_car in lst_cargo else 0
                                                        
                                                        if d_area and d_area not in lst_area: lst_area.append(d_area)
                                                        idx_area = lst_area.index(d_area) if d_area in lst_area else 0

                                                        car = st.selectbox("Cargo", lst_cargo, index=idx_car)
                                                        area_input = st.selectbox("Área", lst_area, index=idx_area) 
                                                        
                                                        rem_b = st.number_input("Remuneración básica", value=d_rem)
                                                        bono = st.text_input("Bonificación", value=d_bon)
                                                        cond = st.text_input("Condición de trabajo", value=d_cond)
                                                        ini = st.date_input("Inicio", value=d_ini, format="DD/MM/YYYY", min_value=date(1950, 1, 1), max_value=date(2100, 12, 31))
                                                        fin = st.date_input("Fin", value=d_fin, format="DD/MM/YYYY", min_value=date(1950, 1, 1), max_value=date(2100, 12, 31))
                                                        t_trab = st.selectbox("Tipo de trabajador", ["Administrativo", "Docente", "Externo"], index=["Administrativo", "Docente", "Externo"].index(d_ttrab))
                                                        mod = st.selectbox("Modalidad", ["Presencial", "Semipresencial", "Virtual"], index=["Presencial", "Semipresencial", "Virtual"].index(d_mod))
                                                        temp = st.selectbox("Temporalidad", ["Plazo fijo", "Plazo indeterminado", "Ordinarizado"], index=["Plazo fijo", "Plazo indeterminado", "Ordinarizado"].index(d_temp))
                                                        lnk = st.text_input("Link")
                                                        tcont = st.selectbox("Tipo Contrato", ["Planilla completo", "Tiempo Parcial", "Recibo por Honorarios", "Otro"], index=["Planilla completo", "Tiempo Parcial", "Recibo por Honorarios", "Otro"].index(d_tcont))
                                                        
                                                        est_a = "ACTIVO" if fin >= date.today() else "CESADO"
                                                        mot_a = st.selectbox("Motivo Cese", ["Vigente"] + MOTIVOS_CESE) if est_a == "CESADO" else "Vigente"

                                                        if st.form_submit_button("Guardar Contrato"):
                                                            nid = dfs[h_name]["id"].max() + 1 if not dfs[h_name].empty else 1
                                                            
                                                            new = {
                                                                "id": nid, 
                                                                "dni": dni_buscado, 
                                                                "cargo": car,
                                                                "area": area_input,
                                                                "remuneracion basica": rem_b,
                                                                "bonificacion": bono,
                                                                "condicion de trabajo": cond,
                                                                "f_inicio": ini,
                                                                "f_fin": fin,
                                                                "tipo de trabajador": t_trab,
                                                                "modalidad": mod,
                                                                "temporalidad": temp,
                                                                "link": lnk,
                                                                "tipo contrato": tcont,
                                                                "estado": est_a,
                                                                "motivo cese": mot_a
                                                            }
                                                            dfs[h_name] = pd.concat([dfs[h_name], pd.DataFrame([new])], ignore_index=True)
                                                            save_data(dfs)
                                                            st.success("✅ Contrato guardado correctamente.")
                                                            st.rerun()
                                                            
                                
                                with col_b:
                                        with st.expander("📝 Editar / Eliminar"):
                                            
                                            # --- INICIO DE LA CORRECCIÓN (BLINDAJE CONTRA NAMEERROR) ---
                                            try:
                                                hay_seleccion = not sel.empty
                                            except NameError:
                                                hay_seleccion = False
                                            # --- FIN DE LA CORRECCIÓN ---
    
                                            if hay_seleccion:
                                                idx = sel.index[0]
                                                with st.form(f"f_edit_{h_name}"):
                                                    if h_name == "CONTRATOS":
                                                        # LEEMOS DE 'sel' EN MAYÚSCULAS (porque así viene de la tabla visual)
                                                        n_car = st.text_input("Cargo", value=str(sel.iloc[0].get("CARGO", "")))
                                                        n_area = st.text_input("Área", value=str(sel.iloc[0].get("AREA", "")))
                                                        
                                                        try: 
                                                            val_rem = float(sel.iloc[0].get("REMUNERACION BASICA", 0.0))
                                                        except: 
                                                            val_rem = 0.0
                                                        n_rem = st.number_input("Remuneración básica", value=val_rem)
                                                        
                                                        n_bon = st.text_input("Bonificación", value=str(sel.iloc[0].get("BONIFICACION", "")))
                                                        n_cond = st.text_input("Condición de trabajo", value=str(sel.iloc[0].get("CONDICION DE TRABAJO", "")))
                                                        
                                                        try: 
                                                            val_ini = sel.iloc[0].get("F_INICIO")
                                                            ini_val = pd.to_datetime(val_ini).date() if pd.notnull(val_ini) else date.today()
                                                        except: 
                                                            val_ini = date.today()
                                                        n_ini = st.date_input("Inicio", value=ini_val, format="DD/MM/YYYY", min_value=date(1950, 1, 1), max_value=date(2100, 12, 31))
                                                        
                                                        try: 
                                                            val_fin = sel.iloc[0].get("F_FIN")
                                                            fin_val = pd.to_datetime(val_fin).date() if pd.notnull(val_fin) else date.today()
                                                        except: 
                                                            fin_val = date.today()
                                                        n_fin = st.date_input("Fin", value=fin_val, format="DD/MM/YYYY", min_value=date(1950, 1, 1), max_value=date(2100, 12, 31))
                                                        
                                                        v_ttrab = str(sel.iloc[0].get("TIPO DE TRABAJADOR", "Administrativo"))
                                                        opts_tt = ["Administrativo", "Docente", "Externo"]
                                                        if v_ttrab not in opts_tt: 
                                                            opts_tt.append(v_ttrab)
                                                        n_ttrab = st.selectbox("Tipo de trabajador", opts_tt, index=opts_tt.index(v_ttrab))
                                                        
                                                        v_mod = str(sel.iloc[0].get("MODALIDAD", "Presencial"))
                                                        opts_mod = ["Presencial", "Semipresencial", "Virtual"]
                                                        if v_mod not in opts_mod: 
                                                            opts_mod.append(v_mod)
                                                        n_mod = st.selectbox("Modalidad", opts_mod, index=opts_mod.index(v_mod))
                                                        
                                                        v_tem = str(sel.iloc[0].get("TEMPORALIDAD", "Plazo fijo"))
                                                        opts_tem = ["Plazo fijo", "Plazo indeterminado", "Ordinarizado"]
                                                        if v_tem not in opts_tem: 
                                                            opts_tem.append(v_tem)
                                                        n_tem = st.selectbox("Temporalidad", opts_tem, index=opts_tem.index(v_tem))
                                                        
                                                        n_lnk = st.text_input("Link", value=str(sel.iloc[0].get("LINK", "")))
                                                        
                                                        v_tcont = str(sel.iloc[0].get("TIPO CONTRATO", "Planilla completo"))
                                                        opts_tcon = ["Planilla completo", "Tiempo Parcial", "Recibo por Honorarios", "Otro"]
                                                        if v_tcont not in opts_tcon: 
                                                            opts_tcon.append(v_tcont)
                                                        n_tcont = st.selectbox("Tipo Contrato", opts_tcon, index=opts_tcon.index(v_tcont))
    
                                                        est_e = "ACTIVO" if n_fin >= date.today() else "CESADO"
                                                        v_mot = str(sel.iloc[0].get("MOTIVO CESE", "Vigente"))
                                                        opts_mot = ["Vigente"] + MOTIVOS_CESE
                                                        if v_mot not in opts_mot: 
                                                            opts_mot.append(v_mot)
                                                        mot_e = st.selectbox("Motivo Cese", opts_mot, index=opts_mot.index(v_mot)) if est_e == "CESADO" else "Vigente"
    
                                                        if st.form_submit_button("Actualizar"):
                                                            # GUARDAMOS EN MINÚSCULAS (Para mantener tu Google Sheet sano y sin duplicados)
                                                            update_vals = {
                                                                "cargo": n_car, 
                                                                "area": n_area,
                                                                "remuneracion basica": n_rem, 
                                                                "bonificacion": n_bon, 
                                                                "condicion de trabajo": n_cond, 
                                                                "f_inicio": n_ini, 
                                                                "f_fin": n_fin, 
                                                                "tipo de trabajador": n_ttrab, 
                                                                "modalidad": n_mod, 
                                                                "temporalidad": n_tem, 
                                                                "link": n_lnk, 
                                                                "tipo contrato": n_tcont, 
                                                                "estado": est_e, 
                                                                "motivo cese": mot_e
                                                            }
                                                            for k, v in update_vals.items(): 
                                                                dfs[h_name].at[idx, k] = v
                                                            save_data(dfs)
                                                            st.rerun()
                                                    else:
                                                        edit_row = {}
                                                        row = sel.iloc[0] 
                                                        
                                                        # ---> NUEVO FILTRO: Eliminamos duplicados y la columna AREA <---
                                                        columnas_limpias = []
                                                        vistas = set()
                                                        for c in cols_reales:
                                                            c_upper = str(c).upper().strip() # Quitamos espacios extra y pasamos a mayúsculas
                                                            # Ignoramos la columna si ya la vimos o si es "AREA"
                                                            if c_upper not in vistas and c_upper != "AREA":
                                                                vistas.add(c_upper)
                                                                columnas_limpias.append(c)
                                                        
                                                        # Ahora iteramos sobre la lista filtrada, no sobre cols_reales
                                                        for i, col in enumerate(columnas_limpias):
                                                            val = row.get(str(col).upper(), "")
                                                            
                                                            if "fecha" in col.lower() or "f_" in col.lower(): 
                                                                if pd.notnull(val) and isinstance(val, (date, datetime)): d_val = val
                                                                elif pd.notnull(val) and isinstance(val, str):
                                                                    try: d_val = pd.to_datetime(val).date()
                                                                    except: d_val = date.today()
                                                                else: d_val = date.today()
                                                                edit_row[col] = st.date_input(col.title(), value=d_val, format="DD/MM/YYYY", min_value=date(1950, 1, 1), max_value=date(2100, 12, 31), key=f"date_{h_name}_{col}_{idx}_{i}")
                                                                
                                                            elif col.lower() == "edad":
                                                                val_edad = int(val) if pd.notnull(val) and str(val).replace('.','',1).isdigit() else 0
                                                                edit_row[col] = st.number_input(col.title(), value=val_edad, disabled=True, key=f"edad_{h_name}_{col}_{idx}_{i}")
                                                                
                                                            elif col.lower() in ["remuneración", "bonificación", "sueldo", "días generados", "dias gozados", "saldo", "monto", "remuneracion basica", "bonificacion"]: 
                                                                try: n_val = float(val) if pd.notnull(val) else 0.0
                                                                except: n_val = 0.0
                                                                edit_row[col] = st.number_input(col.title(), value=n_val, key=f"num_{h_name}_{col}_{idx}_{i}")
                                                                
                                                            else: 
                                                                edit_row[col] = st.text_input(col.title(), value=str(val) if pd.notnull(val) else "", key=f"text_{h_name}_{col}_{idx}_{i}")
    
                                                        col_btn1, col_btn2 = st.columns(2)
                                                        with col_btn1:
                                                            if st.form_submit_button("Actualizar Registro"):
                                                                for k, v in edit_row.items(): 
                                                                    dfs[h_name].at[idx, k] = v
                                                                save_data(dfs)
                                                                st.rerun()
                                                        with col_btn2:
                                                            if st.form_submit_button("🗑️ Eliminar Registro", type="primary"):
                                                                dfs[h_name] = dfs[h_name].drop(idx)
                                                                save_data(dfs)
                                                                st.rerun()
                                            else:
                                                st.info("Activa la casilla (SEL) en la tabla superior para editar o eliminar el registro. Si estás buscando a alguien que ingresaste a mano, verifica que su nombre en el Google Sheets no tenga espacios al final.")
                
    # ==========================================
    # --- SECCIÓN REGISTRO Y NÓMINA ---
    # ==========================================
    elif m == "➕ Registro" and not es_lector:
        import mod_registro
        mod_registro.mostrar(dfs, save_data)
        
    # (El resto del código hacia abajo queda igualito)

    elif m == "📊 Nómina General":
        mod_nomina.mostrar(dfs, save_data)
    # ==========================================
    # MÓDULO: ESTRUCTURA Y PUESTOS (MOF)
    # ==========================================
    elif m == "🏢 Estructura":
        mod_estructura.mostrar(dfs)    
    
    # ==========================================
    # MÓDULO: REPORTE GENERAL
    # ==========================================
    elif m == "Reporte General": # (Asegúrate de que 'm' sea la variable de tu menú)
        mod_reportegeneral.mostrar(dfs)

    # ==========================================
    # MÓDULO: REPORTE DE SALDO DE VACACIONES
    # ==========================================
    elif m == "Vacaciones":
        mod_vacaciones.mostrar(dfs)

    # ==========================================
    # MÓDULO: CUMPLEAÑEROS
    # ==========================================
    elif m == "Cumpleañeros":
        mod_cumpleanos.mostrar(dfs)   
    
    # ==========================================
    # MÓDULO: VENCIMIENTO DE CONTRATOS
    # ==========================================
    elif m == "Vencimientos":
        mod_vencimientos.mostrar(dfs)

    # ==========================================
    # MÓDULO: DASHBOARD DE DESEMPEÑO
    # ==========================================
    elif m == "📈 Dashboard Desempeño":
        import mod_reportes
        mod_reportes.mostrar(dfs)

    # ==========================================
    # MÓDULO: GESTOR DE EVALUACIONES
    # ==========================================
    # ... dentro de tu menú en app.py ...
    if m == "📋 Evaluaciones":  # (o como se llame tu opción en el menú)
        mod_gestor_evaluaciones.mostrar(dfs, save_data)
